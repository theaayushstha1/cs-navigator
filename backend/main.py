import sys
# Force unbuffered output so we see logs immediately
sys.stdout.reconfigure(line_buffering=True)

print("[OK] main.py loaded successfully")

import os
import re
import json
import asyncio
# import time  # Commented: currently unused, kept for potential future use
import shutil #  NEW: For file operations
from typing import List, Dict, Any, Optional
from contextlib import asynccontextmanager
from datetime import datetime, timezone, timedelta

#  FIXED IMPORTS: Use 'pypdf' which you installed, not 'PyPDF2'
import pypdf 
import docx
from langchain.schema import SystemMessage, HumanMessage 

from fastapi import FastAPI, HTTPException, Depends, status, File, UploadFile, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.trustedhost import TrustedHostMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse, HTMLResponse
from pydantic import BaseModel, field_validator
from collections import Counter
import io
from dotenv import load_dotenv

# ==============================================================================
# 1. ENVIRONMENT LOADING (FIXED FOR ROOT FOLDER)
# ==============================================================================
# Get the absolute path of the backend folder
BACKEND_DIR = os.path.dirname(os.path.abspath(__file__))
# Get the project root (one level up)
PROJECT_ROOT = os.path.dirname(BACKEND_DIR)
# Path to .env file in the root
ENV_PATH = os.path.join(PROJECT_ROOT, ".env")

# Load course catalog for context injection
COURSE_CATALOG_TEXT = ""
_catalog_path = os.path.join(BACKEND_DIR, "data_sources", "classes.json")
if os.path.exists(_catalog_path):
    try:
        with open(_catalog_path) as _f:
            _catalog = json.load(_f)
        _lines = []
        for c in _catalog.get("courses", []):
            prereqs = ", ".join(c.get("prerequisites", [])) or "None"
            _lines.append(f"  {c['course_code']} - {c['course_name']} ({c.get('credits',3)} cr, {c.get('category','')}) Prereqs: {prereqs}")
        COURSE_CATALOG_TEXT = "AVAILABLE CS COURSES AT MORGAN STATE (from official catalog):\n" + "\n".join(_lines) + "\n"
    except Exception as _e:
        print(f"[WARN] Failed to load course catalog: {_e}")

print(f"[INFO] Looking for .env at: {ENV_PATH}")

if os.path.exists(ENV_PATH):
    load_dotenv(ENV_PATH)
    print("[OK] .env file loaded!")
else:
    print("[ERROR] .env file NOT found at root. Checking backend folder...")
    load_dotenv(os.path.join(BACKEND_DIR, ".env"))

print(f"[KEY] JWT_SECRET Check: {'FOUND' if os.getenv('JWT_SECRET') else 'MISSING'}")

# SQLAlchemy Imports
from sqlalchemy.orm import Session
from sqlalchemy.exc import OperationalError, ProgrammingError
from sqlalchemy import Column, Integer, String, Text, DateTime, ForeignKey, text

# Vertex AI Agent Engine (replaces Pinecone + OpenAI RAG pipeline)
from vertex_agent import query_agent, query_agent_stream, check_agent_health, reset_session

# Query caching for faster responses
from cache import query_cache, get_context_hash, log_cache_stats


# Legacy imports kept for /ingest endpoint and file analysis fallback
try:
    from langchain.text_splitter import TokenTextSplitter
    from langchain_openai import OpenAIEmbeddings
    from langchain_pinecone import PineconeVectorStore
    from langchain_community.chat_models import ChatOpenAI
    from langchain.chains import RetrievalQA
    from pinecone import Pinecone
    LEGACY_RAG_AVAILABLE = True
except ImportError:
    LEGACY_RAG_AVAILABLE = False
    print("   Legacy RAG imports not available (Pinecone/LangChain not installed)")

# Local Imports (Auth & DB) - These must run AFTER load_dotenv
from db import SessionLocal, engine, Base
from models import User, DegreeWorksData, BannerStudentData, SupportTicket, FailedQuery, KBSuggestion, CanvasStudentData, UserMemory, ChatHistory, Feedback
from security import hash_password, verify_password, create_access_token
from jose import JWTError, jwt

# Banner SSB integration (CAS auth + REST API sync)
from banner_scraper import sync_banner

# ==============================================================================
# 2. CONFIGURATION & CONSTANTS
# ==============================================================================
# Banner sync rate limiting: {user_id: [timestamp, ...]}
_banner_sync_timestamps: dict[int, list] = {}
# Vertex AI Agent Engine config
USE_VERTEX_AGENT   = os.getenv("USE_VERTEX_AGENT", "true").lower() == "true"
ADK_BASE_URL       = os.getenv("ADK_BASE_URL", "http://127.0.0.1:8080")

# Legacy Pinecone + OpenAI config (kept for /ingest and TTS)
PINECONE_API_KEY   = os.getenv("PINECONE_API_KEY")
PINECONE_ENV       = os.getenv("PINECONE_ENV")
PINECONE_INDEX     = os.getenv("PINECONE_INDEX_NAME")
PINECONE_NAMESPACE = os.getenv("PINECONE_NAMESPACE", "docs")
OPENAI_API_KEY     = os.getenv("OPENAI_API_KEY")  # Still needed for TTS
JWT_SECRET         = os.getenv("JWT_SECRET")
ALGORITHM          = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES", "4320"))  # 3 days default

# Upload configuration
UPLOAD_FOLDER = os.path.join(BACKEND_DIR, "uploads", "profile_pictures")
CHAT_FILES_FOLDER = os.path.join(BACKEND_DIR, "uploads", "chat_files") #  NEW: Chat files folder

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'pdf', 'txt', 'docx', 'doc', 'mov', 'mp4'} #  NEW: Added Docs

# Create folders if not exist
for folder in [UPLOAD_FOLDER, CHAT_FILES_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder, exist_ok=True)
        print(f"[OK] Created folder: {folder}")

# Safety check for keys
if USE_VERTEX_AGENT:
    print(f"[INFO] Using Vertex AI Agent Engine at {ADK_BASE_URL}")
elif not all([PINECONE_API_KEY, PINECONE_ENV, PINECONE_INDEX, OPENAI_API_KEY]):
    print("[WARN] WARNING: Some API keys are missing. Chatbot features will be limited.")

# ==============================================================================
# 3. DATABASE MODELS
# ==============================================================================
# ChatHistory, Feedback, and all other models are now in models.py
# Imported above: ChatHistory, Feedback (via models import line)

def init_db():
    """Initializes the database tables and runs migrations."""
    # 1. Create tables if missing
    try:
        Base.metadata.create_all(bind=engine)
        print("[OK] Database tables checked/created.")
    except Exception as e:
        print(f"[WARN] DB Connection Error: {e}")

    # 2. Add session_id column if missing (For existing DBs)
    with engine.connect() as conn:
        try:
            # Check if column exists by selecting from it
            conn.execute(text("SELECT session_id FROM chat_history LIMIT 1"))
        except (OperationalError, ProgrammingError):
            print("[WARN] 'session_id' column missing. Adding it now...")
            try:
                conn.execute(text("ALTER TABLE chat_history ADD COLUMN session_id VARCHAR(255) DEFAULT 'default'"))
                conn.commit()
                print("[OK] Successfully added 'session_id' column!")
            except Exception as e:
                print(f"[ERROR] Failed to add column: {e}")

        # 3. Add profile_picture_data column if missing (For base64 storage)
        try:
            conn.execute(text("SELECT profile_picture_data FROM users LIMIT 1"))
        except (OperationalError, ProgrammingError):
            print("[WARN] 'profile_picture_data' column missing. Adding it now...")
            try:
                conn.execute(text("ALTER TABLE users ADD COLUMN profile_picture_data LONGTEXT"))
                conn.commit()
                print("[OK] Successfully added 'profile_picture_data' column!")
            except Exception as e:
                print(f"[ERROR] Failed to add profile_picture_data column: {e}")

        # 4. Add morgan_connected_at column if missing
        try:
            conn.execute(text("SELECT morgan_connected_at FROM users LIMIT 1"))
        except (OperationalError, ProgrammingError):
            print("[WARN] 'morgan_connected_at' column missing. Adding it now...")
            try:
                conn.execute(text("ALTER TABLE users ADD COLUMN morgan_connected_at DATETIME"))
                conn.commit()
                print("[OK] Successfully added 'morgan_connected_at' column!")
            except Exception as e:
                print(f"[ERROR] Failed to add morgan_connected_at column: {e}")

        # 5. Add email auth columns if missing
        for col, col_type in [
            ("email_verified", "BOOLEAN DEFAULT TRUE"),
            ("verification_token", "VARCHAR(255)"),
            ("reset_token", "VARCHAR(255)"),
            ("reset_token_expires", "DATETIME"),
        ]:
            try:
                conn.execute(text(f"SELECT {col} FROM users LIMIT 1"))
            except (OperationalError, ProgrammingError):
                try:
                    conn.execute(text(f"ALTER TABLE users ADD COLUMN {col} {col_type}"))
                    conn.commit()
                    print(f"[OK] Added '{col}' column to users")
                except Exception:
                    pass

        # 6. Check if degreeworks_data table exists
        try:
            conn.execute(text("SELECT id FROM degreeworks_data LIMIT 1"))
            print("[OK] degreeworks_data table exists")
        except (OperationalError, ProgrammingError):
            print("[WARN] 'degreeworks_data' table missing. Creating it now...")
            try:
                conn.execute(text("""
                    CREATE TABLE degreeworks_data (
                        id INT AUTO_INCREMENT PRIMARY KEY,
                        user_id INT UNIQUE NOT NULL,
                        student_name VARCHAR(255),
                        student_id VARCHAR(50),
                        degree_program VARCHAR(255),
                        catalog_year VARCHAR(20),
                        classification VARCHAR(50),
                        advisor VARCHAR(255),
                        overall_gpa FLOAT,
                        major_gpa FLOAT,
                        total_credits_earned FLOAT,
                        credits_required FLOAT,
                        credits_remaining FLOAT,
                        courses_completed TEXT,
                        courses_in_progress TEXT,
                        courses_remaining TEXT,
                        requirements_status TEXT,
                        raw_data TEXT,
                        synced_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                        updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
                    )
                """))
                conn.commit()
                print("[OK] Successfully created 'degreeworks_data' table!")
            except Exception as e:
                print(f"[ERROR] Failed to create degreeworks_data table: {e}")

        # 6b. Add data_source column to degreeworks_data if missing
        try:
            conn.execute(text("SELECT data_source FROM degreeworks_data LIMIT 1"))
        except (OperationalError, ProgrammingError):
            print("[WARN] 'data_source' column missing from degreeworks_data. Adding it now...")
            try:
                conn.execute(text("ALTER TABLE degreeworks_data ADD COLUMN data_source VARCHAR(50) DEFAULT 'manual_entry'"))
                conn.commit()
                print("[OK] Successfully added 'data_source' column!")
            except Exception as e:
                print(f"[ERROR] Failed to add data_source column: {e}")

        # 6. Check if support_tickets table exists
        try:
            conn.execute(text("SELECT id FROM support_tickets LIMIT 1"))
            print("[OK] support_tickets table exists")
        except (OperationalError, ProgrammingError):
            print("[WARN] 'support_tickets' table missing. Creating it now...")
            try:
                conn.execute(text("""
                    CREATE TABLE support_tickets (
                        id INT AUTO_INCREMENT PRIMARY KEY,
                        user_id INT NOT NULL,
                        subject VARCHAR(255) NOT NULL,
                        category VARCHAR(50) NOT NULL,
                        description TEXT NOT NULL,
                        attachment_data LONGTEXT,
                        attachment_name VARCHAR(255),
                        status VARCHAR(50) DEFAULT 'open',
                        priority VARCHAR(20) DEFAULT 'normal',
                        admin_notes TEXT,
                        resolved_by INT,
                        resolved_at DATETIME,
                        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                        updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
                        FOREIGN KEY (resolved_by) REFERENCES users(id) ON DELETE SET NULL
                    )
                """))
                conn.commit()
                print("[OK] Successfully created 'support_tickets' table!")
            except Exception as e:
                print(f"[ERROR] Failed to create support_tickets table: {e}")

        # 7. Check if banner_student_data table exists
        try:
            conn.execute(text("SELECT id FROM banner_student_data LIMIT 1"))
            print("[OK] banner_student_data table exists")
        except (OperationalError, ProgrammingError):
            print("[WARN] 'banner_student_data' table missing. Creating it now...")
            try:
                conn.execute(text("""
                    CREATE TABLE banner_student_data (
                        id INT AUTO_INCREMENT PRIMARY KEY,
                        user_id INT UNIQUE NOT NULL,
                        student_phone VARCHAR(20),
                        student_address TEXT,
                        current_term VARCHAR(50),
                        registered_courses TEXT,
                        total_registered_credits FLOAT,
                        registration_history TEXT,
                        grade_history TEXT,
                        cumulative_gpa FLOAT,
                        total_credits_earned FLOAT,
                        total_credits_attempted FLOAT,
                        deans_list_terms TEXT,
                        synced_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                        updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
                    )
                """))
                conn.commit()
                print("[OK] Successfully created 'banner_student_data' table!")
            except Exception as e:
                print(f"[ERROR] Failed to create banner_student_data table: {e}")

    # 8. Create/Update admin account
    try:
        db = SessionLocal()
        admin_email = os.getenv("ADMIN_EMAIL", "admin@morgan.edu")
        admin_password = os.getenv("ADMIN_PASSWORD")
        if not admin_password:
            print("[WARN] ADMIN_PASSWORD not set in env, skipping admin account creation")
            db.close()
            return

        existing_admin = db.query(User).filter(User.email == admin_email).first()

        if existing_admin:
            # Update existing user to admin
            if existing_admin.role != "admin":
                existing_admin.role = "admin"
                db.commit()
                print(f"[OK] Updated {admin_email} to admin role!")
            else:
                print(f"[OK] Admin account {admin_email} already exists with admin role.")
        else:
            # Create new admin account
            from security import hash_password
            hashed = hash_password(admin_password)
            admin_user = User(
                email=admin_email,
                password_hash=hashed,
                role="admin",
                name="Admin"
            )
            db.add(admin_user)
            db.commit()
            print(f"[OK] Created admin account: {admin_email}")

        db.close()
    except Exception as e:
        print(f"[ERROR] Failed to create/update admin account: {e}")

init_db()

# ==============================================================================
# 4. FASTAPI APP SETUP
# ==============================================================================
# AI System globals (initialized in lifespan)
pc = None
retriever = None
qa = None
llm = None

def build_qa_chain():
    """Initialize legacy AI components on startup (only when not using Vertex AI)"""
    global retriever, qa, llm, pc
    if USE_VERTEX_AGENT:
        # Check Vertex AI Agent health
        health = check_agent_health()
        print(f" Vertex AI Agent: {health['status']} - {health['message']}")
        if health["status"] != "connected":
            print("[WARN] ADK server not running. Start it with:")
            print("   cd google-ai-engine-research/adk_deploy && python -m google.adk.cli web . --port 8080")
        return

    if not LEGACY_RAG_AVAILABLE:
        print("[WARN] Legacy RAG libraries not installed. Chatbot will be offline.")
        return
    if not all([PINECONE_API_KEY, OPENAI_API_KEY, PINECONE_INDEX]):
        print("[WARN] API Keys missing. Chatbot will be offline.")
        return
    try:
        pc = Pinecone(api_key=PINECONE_API_KEY)
        embeddings = OpenAIEmbeddings(model="text-embedding-3-small", openai_api_key=OPENAI_API_KEY)
        store = PineconeVectorStore.from_existing_index(
            index_name=PINECONE_INDEX,
            embedding=embeddings,
            namespace=PINECONE_NAMESPACE,
        )
        retriever = store.as_retriever(
            search_type="mmr",
            search_kwargs={
                "k": 10,
                "fetch_k": 30,
                "lambda_mult": 0.5
            }
        )
        llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, model_name="gpt-3.5-turbo", temperature=0)
        qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever, return_source_documents=True)
        print("[OK] Legacy AI System Initialized (Pinecone + OpenAI)")
    except Exception as e:
        print(f"[ERROR] AI Init Failed: {e}")

@asynccontextmanager
async def lifespan(app):
    """Modern lifespan event handler for FastAPI"""
    # Startup
    build_qa_chain()
    yield
    # Shutdown (cleanup if needed)

app = FastAPI(title="CS Navigator API", version="5.0.0", lifespan=lifespan)

ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://localhost:5173,http://localhost:5174,http://localhost:5175,http://127.0.0.1:3000,http://127.0.0.1:5173,http://127.0.0.1:5174,http://127.0.0.1:8000,https://inavigator.ai,https://cs.inavigator.ai,https://api.inavigator.ai,https://csnavigator-frontend-750361124802.us-central1.run.app").split(",")
print(f"[CORS] Allowed origins: {ALLOWED_ORIGINS}")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.add_middleware(
    TrustedHostMiddleware,
    allowed_hosts=os.getenv("TRUSTED_HOSTS", "localhost,127.0.0.1,inavigator.ai,cs.inavigator.ai,api.inavigator.ai,csnavigator-backend-750361124802.us-central1.run.app,csnavigator-frontend-750361124802.us-central1.run.app,csnavigator-backend-jvat5svbjq-uc.a.run.app").split(",")
)

# Mount Static Files (Profile Pictures AND Chat Files)
UPLOADS_DIR = os.path.join(BACKEND_DIR, "uploads")
if os.path.exists(UPLOADS_DIR):
    try:
        app.mount("/uploads", StaticFiles(directory=UPLOADS_DIR), name="uploads")
        print(f"[OK] Static files mounted: /uploads -> {UPLOADS_DIR}")
    except Exception as e:
        print(f"[ERROR] Error mounting static files: {e}")
else:
    os.makedirs(UPLOADS_DIR, exist_ok=True)
    print(f"[OK] Created uploads directory: {UPLOADS_DIR}")

# ==============================================================================
# 5. AUTHENTICATION HELPERS
# ==============================================================================
security = HTTPBearer()

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def get_current_user(
    credentials: HTTPAuthorizationCredentials = Depends(security),
    db: Session = Depends(get_db)
) -> Dict[str,Any]:
    token = credentials.credentials
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[ALGORITHM])
        user_email = payload.get("email")
        if not user_email:
            raise HTTPException(status_code=403, detail="Invalid token")

        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            raise HTTPException(status_code=403, detail="User not found")

        return {
            "user_id": user.id,
            "email": user.email,
            "role": user.role
        }
    except JWTError as e:
        print(f"JWT decode error: {e}")
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Invalid token")

def allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# ==============================================================================
# 6. PYDANTIC SCHEMAS
# ==============================================================================
class RegisterRequest(BaseModel):
    email: str
    password: str

    @staticmethod
    def validate_email_format(v):
        import re
        if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', v):
            raise ValueError("Invalid email format")
        return v

    @staticmethod
    def validate_password_strength(v):
        if len(v) < 8:
            raise ValueError("Password must be at least 8 characters")
        return v

class LoginRequest(BaseModel):
    email: str
    password: str

VALID_MODELS = {"", "inav-1.0", "inav-1.1"}

class QueryRequest(BaseModel):
    query: str
    session_id: str = "default"
    skip_cache: bool = False
    model: str = ""              # "inav-1.0" (fast) or "inav-1.1" (pro)

    @field_validator("model", mode="before")
    @classmethod
    def validate_model(cls, v):
        if v not in VALID_MODELS:
            return ""
        return v

class GuestQueryRequest(BaseModel):
    query: str
    guestProfile: Optional[dict] = None

# ==============================================================================
# GUEST RATE LIMITING (Simple In-Memory)
# ==============================================================================
from collections import defaultdict
import time as time_module

guest_rate_limits = defaultdict(list)  # IP -> list of timestamps
GUEST_RATE_LIMIT = 15  # requests per minute (time-based session provides natural limiting)
GUEST_RATE_WINDOW = 60  # seconds
_guest_rate_last_cleanup = time_module.time()

def check_guest_rate_limit(ip: str) -> bool:
    """Check if IP is within rate limit. Returns True if allowed, False if blocked."""
    global _guest_rate_last_cleanup
    current_time = time_module.time()

    # Periodic cleanup: purge stale IPs every 10 minutes to prevent memory leak
    if current_time - _guest_rate_last_cleanup > 600:
        stale_ips = [k for k, v in guest_rate_limits.items() if not v or current_time - v[-1] > GUEST_RATE_WINDOW]
        for k in stale_ips:
            del guest_rate_limits[k]
        _guest_rate_last_cleanup = current_time

    # Clean old entries for this IP
    guest_rate_limits[ip] = [t for t in guest_rate_limits[ip] if current_time - t < GUEST_RATE_WINDOW]
    # Check limit
    if len(guest_rate_limits[ip]) >= GUEST_RATE_LIMIT:
        return False
    # Add new request
    guest_rate_limits[ip].append(current_time)
    return True

# Forgot-password rate limiting: {email: [timestamp, ...]}
_forgot_pw_timestamps: dict[str, list] = {}
_forgot_pw_last_cleanup = time_module.time()
FORGOT_PW_RATE_LIMIT = 5   # max requests per window
FORGOT_PW_RATE_WINDOW = 900  # 15 minutes

class Course(BaseModel):
    course_code: str
    course_name: str
    credits: int
    prerequisites: List[str] = []
    offered: List[str] = []

class ProfileUpdateRequest(BaseModel):
    name: Optional[str] = None
    studentId: Optional[str] = None
    major: Optional[str] = None

class PasswordChangeRequest(BaseModel):
    currentPassword: str
    newPassword: str

class TTSRequest(BaseModel):
    text: str
    voice: str = "alloy"  # Options: alloy, echo, fable, onyx, nova, shimmer

#  DegreeWorks Data Schema
class DegreeWorksRequest(BaseModel):
    student_name: Optional[str] = None
    student_id: Optional[str] = None
    degree_program: Optional[str] = None
    catalog_year: Optional[str] = None
    classification: Optional[str] = None
    advisor: Optional[str] = None
    overall_gpa: Optional[float] = None
    major_gpa: Optional[float] = None
    total_credits_earned: Optional[float] = None
    credits_required: Optional[float] = None
    credits_remaining: Optional[float] = None
    courses_completed: Optional[List[Dict[str, Any]]] = None  # [{code, name, credits, grade, semester}]
    courses_in_progress: Optional[List[Dict[str, Any]]] = None  # [{code, name, credits, semester}]
    courses_remaining: Optional[List[Dict[str, Any]]] = None  # [{code, name, credits, category}]
    requirements_status: Optional[List[Dict[str, Any]]] = None  # [{category, status, details}]
    raw_data: Optional[str] = None

# ==============================================================================
# 7. STATIC DATA & RESOURCES
# ==============================================================================
DATA_DIR       = os.path.join(BACKEND_DIR, "data_sources")
CLASSES_FILE   = os.path.join(DATA_DIR, "classes.json")
KB_COURSES_FILE = os.path.join(DATA_DIR, "courses.txt")
RESOURCES_FILE = os.path.join(DATA_DIR, "academic_resources.json")

# Cached parsed curriculum from txt source of truth
_parsed_curriculum = None

def parse_curriculum_from_txt():
    """Parse courses.txt into the structured JSON format the frontend expects.
    This makes the txt knowledge base files the single source of truth for the curriculum page."""
    global _parsed_curriculum
    if _parsed_curriculum is not None:
        return _parsed_curriculum

    degree_info = {
        "program": "Computer Science, B.S.",
        "university": "Morgan State University",
        "total_credits": 120,
        "general_education_credits": 44,
        "supporting_credits": 11,
        "major_credits": 65,
        "cs_core_credits": 76,
        "description": "A minimum of 120 credit hours are required to graduate with a B.S. in Computer Science."
    }

    elective_requirements = {
        "group_a": {"name": "Group A Electives", "required_courses": 3,
                    "description": "Students must choose three (3) courses from Group A"},
        "group_b": {"name": "Group B Electives", "required_courses": 2,
                    "description": "Students must choose two (2) courses from Group B"},
        "group_c": {"name": "Group C Electives", "required_courses": 4,
                    "description": "Students must choose four (4) courses from Group C. Note: COSC 470 OR COSC 472 - only one counts."},
        "group_d": {"name": "Group D Electives", "required_courses": 1,
                    "description": "Students must choose one (1) course from Group D, or any 300-400 level COSC course not previously taken"}
    }

    section_map = {
        "REQUIRED COURSES": ("Required", "required", None),
        "SUPPORTING COURSES": ("Supporting", "supporting", None),
        "GROUP A ELECTIVES": ("Group A Elective", "group_a", "Choose 3 courses from Group A"),
        "GROUP B ELECTIVES": ("Group B Elective", "group_b", "Choose 2 courses from Group B"),
        "GROUP C ELECTIVES": ("Group C Elective", "group_c", "Choose 4 courses from Group C (COSC 470 OR COSC 472)"),
        "GROUP D ELECTIVES": ("Group D Elective", "group_d", "Choose 1 course from Group D"),
    }

    courses = []
    with open(KB_COURSES_FILE, encoding="utf-8") as f:
        lines = f.read().split('\n')

    current_cat = current_req = current_note = None
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Detect section headers
        matched = False
        for key, (cat, req, note) in section_map.items():
            if line.upper().startswith(key):
                current_cat, current_req, current_note = cat, req, note
                matched = True
                break
        if matched:
            i += 1
            continue

        # Detect course line: "COSC 111 - Introduction to Computer Science I"
        m = re.match(r'^([A-Z]+\s+\d{3})\s*[-\u2013\u2014]\s*(.+)$', line)
        if m and current_cat:
            course = {
                "course_code": m.group(1).strip(),
                "course_name": m.group(2).strip(),
                "credits": 3,
                "category": current_cat,
                "requirement_type": current_req,
                "prerequisites": [],
                "offered": ["Fall", "Spring"],
            }
            if current_note:
                course["elective_note"] = current_note

            # Parse detail lines until blank line
            i += 1
            while i < len(lines) and lines[i].strip():
                d = lines[i].strip()
                if d.lower().startswith("credits:"):
                    try:
                        course["credits"] = int(d.split(":", 1)[1].strip())
                    except ValueError:
                        pass
                elif d.lower().startswith("prerequisite"):
                    raw = d.split(":", 1)[1].strip()
                    if raw.lower() in ("none", ""):
                        course["prerequisites"] = []
                    else:
                        parts = [p.strip() for p in raw.split(",")]
                        course["prerequisites"] = [
                            p[3:].strip() if p.startswith("or ") else p
                            for p in parts if p
                        ]
                elif d.lower().startswith("offered:"):
                    course["offered"] = [o.strip() for o in d.split(":", 1)[1].split(",") if o.strip()]
                elif d.lower().startswith("also satisfies"):
                    course["note"] = d
                i += 1

            courses.append(course)
            continue

        i += 1

    result = {
        "degree_info": degree_info,
        "courses": courses,
        "elective_requirements": elective_requirements
    }
    _parsed_curriculum = result
    return result

helpful_links = {}
if os.path.exists(RESOURCES_FILE):
    try:
        with open(RESOURCES_FILE, "r", encoding="utf-8") as f:
            res_data = json.load(f)
        helpful_links = res_data.get("academic_and_student_support", {}).get("helpful_links", {})
    except:
        pass

def load_json_documents(paths: List[str]) -> List[Dict[str,Any]]:
    docs: List[Dict[str,Any]] = []
    for p in paths:
        try:
            data = json.load(open(p, encoding="utf-8"))
            if isinstance(data, dict):
                for k, v in data.items():
                    if isinstance(v, dict):
                        parts = [f"{subk}: {subv}" for subk, subv in v.items()]
                        docs.append({"text": f"{k} – " + "; ".join(parts), "source": p})
                    else:
                        docs.append({"text": f"{k}: {v}", "source": p})
            elif isinstance(data, list):
                for obj in data:
                    text = "\n".join(f"{kk}: {vv}" for kk, vv in obj.items())
                    docs.append({"text": text, "source": p})
        except Exception:
            pass
    return docs

# ==============================================================================
# 7b. ROOT DASHBOARD - Show endpoints & recent logs
# ==============================================================================
import logging
from collections import deque

# In-memory log buffer (last 200 log lines)
_log_buffer = deque(maxlen=200)

class BufferHandler(logging.Handler):
    def emit(self, record):
        _log_buffer.append(self.format(record))

_buf_handler = BufferHandler()
_buf_handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(name)s: %(message)s"))
logging.getLogger().addHandler(_buf_handler)
logging.getLogger("uvicorn.access").addHandler(_buf_handler)
logging.getLogger("uvicorn.error").addHandler(_buf_handler)

def get_optional_user(
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(HTTPBearer(auto_error=False)),
    db: Session = Depends(get_db)
) -> Optional[Dict[str, Any]]:
    """Like get_current_user but returns None instead of 401/403 when unauthenticated."""
    if not credentials:
        return None
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=[ALGORITHM])
        user_email = payload.get("email")
        if not user_email:
            return None
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            return None
        return {"user_id": user.id, "email": user.email, "role": user.role}
    except JWTError:
        return None

@app.get("/", response_class=HTMLResponse)
def root_dashboard(request: Request, user: Optional[dict] = Depends(get_optional_user)):
    """Dashboard showing all endpoints and recent logs. Admin only, dev/staging only."""
    if not user or user.get("role") != "admin":
        from starlette.responses import RedirectResponse
        return RedirectResponse(url="/docs")
    # Hide logs in production unless explicitly enabled
    show_logs = os.getenv("SHOW_DASHBOARD_LOGS", "true").lower() == "true"
    routes = []
    for route in request.app.routes:
        if hasattr(route, "methods"):
            for method in sorted(route.methods):
                if method == "HEAD":
                    continue
                routes.append({"method": method, "path": route.path})
    routes.sort(key=lambda r: (r["path"], r["method"]))

    import html as _html
    logs_html = "\n".join(
        f"<div class='log'>{_html.escape(line)}</div>" for line in reversed(_log_buffer)
    ) or "<div class='log dim'>No logs captured yet.</div>"

    rows = "\n".join(
        f"<tr><td class='method {r['method'].lower()}'>{r['method']}</td><td>{r['path']}</td></tr>"
        for r in routes
    )

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>CSNavigator API</title>
<style>
  * {{ margin:0; padding:0; box-sizing:border-box; }}
  body {{ font-family: 'SF Mono', 'Fira Code', monospace; background:#0d1117; color:#c9d1d9; padding:2rem; }}
  h1 {{ color:#58a6ff; margin-bottom:.5rem; font-size:1.4rem; }}
  h2 {{ color:#8b949e; margin:1.5rem 0 .5rem; font-size:1rem; text-transform:uppercase; letter-spacing:.1em; }}
  .info {{ color:#8b949e; font-size:.85rem; margin-bottom:1rem; }}
  table {{ border-collapse:collapse; width:100%; max-width:700px; }}
  td {{ padding:4px 12px; border-bottom:1px solid #21262d; font-size:.85rem; }}
  .method {{ font-weight:bold; width:60px; }}
  .get {{ color:#3fb950; }}  .post {{ color:#d29922; }}  .put {{ color:#58a6ff; }}  .delete {{ color:#f85149; }}
  #logs {{ background:#161b22; border:1px solid #30363d; border-radius:6px; padding:1rem; max-height:500px; overflow-y:auto; margin-top:.5rem; }}
  .log {{ font-size:.78rem; padding:2px 0; border-bottom:1px solid #21262d; white-space:pre-wrap; word-break:break-all; }}
  .dim {{ color:#484f58; }}
  .refresh {{ color:#58a6ff; text-decoration:none; font-size:.85rem; }}
</style></head><body>
  <h1>CSNavigator API v2.1.0</h1>
  <div class="info">Backend is running. {len(routes)} endpoints registered.</div>

  <h2>Endpoints</h2>
  <table>{rows}</table>

  {'<h2>Recent Logs <a class="refresh" href="/">refresh</a></h2><div id="logs">' + logs_html + '</div>' if show_logs else '<p class="dim">Logs hidden in production. Set SHOW_DASHBOARD_LOGS=true to enable.</p>'}
</body></html>"""

# ==============================================================================
# 8. API ENDPOINTS
# ==============================================================================

# --- Auth ---
ALLOWED_EMAIL_DOMAINS = ["morgan.edu"]

_register_timestamps: dict[str, list] = {}

@app.post("/api/register", status_code=status.HTTP_201_CREATED)
def register(req: RegisterRequest, db: Session = Depends(get_db)):
    import re
    from email_service import generate_token, send_verification_email

    email = req.email.strip().lower()

    if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email):
        raise HTTPException(status_code=400, detail="Invalid email format")

    # Rate limit per EMAIL (not per IP). On campus WiFi all students share one IP,
    # so IP-based limiting blocks innocent users. 3 attempts per email per hour.
    now_ts = time_module.time()
    reg_ts = _register_timestamps.get(email, [])
    reg_ts = [t for t in reg_ts if now_ts - t < 3600]
    if len(reg_ts) >= 3:
        raise HTTPException(status_code=429, detail="Too many attempts for this email. Try again in an hour.")
    reg_ts.append(now_ts)
    _register_timestamps[email] = reg_ts

    # Only allow Morgan State email for new registrations
    email_domain = email.split("@")[-1].lower()
    allow_test = os.getenv("ALLOW_TEST_EMAILS", "false").lower() == "true"
    if email_domain not in ALLOWED_EMAIL_DOMAINS and not (allow_test and email.endswith("@test.com")):
        raise HTTPException(status_code=400, detail="Only @morgan.edu email addresses are allowed.")

    if len(req.password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters")
    if db.query(User).filter(User.email == req.email).first():
        raise HTTPException(status_code=400, detail="Email already registered")

    hashed = hash_password(req.password)
    token = generate_token()
    student = User(email=req.email, password_hash=hashed, role="student", email_verified=False, verification_token=token)
    db.add(student)
    db.commit()
    db.refresh(student)

    send_verification_email(req.email, token)
    return {"message": "Account created! Check your Morgan State email to verify.", "user_id": student.id}


@app.get("/api/verify-email")
def verify_email(token: str, db: Session = Depends(get_db)):
    from starlette.responses import RedirectResponse
    user = db.query(User).filter(User.verification_token == token).first()
    if not user:
        raise HTTPException(status_code=400, detail="Invalid or expired verification link")
    user.email_verified = True
    user.verification_token = None
    db.commit()
    # Redirect to login with success flag
    app_url = os.getenv("APP_URL", "https://cs.inavigator.ai")
    return RedirectResponse(url=f"{app_url}/login?verified=true")


@app.post("/api/resend-verification")
async def resend_verification(request: Request, db: Session = Depends(get_db)):
    from email_service import generate_token, send_verification_email
    body = await request.json()
    email = body.get("email", "")
    user = db.query(User).filter(User.email == email).first()
    if not user:
        return {"message": "If an account exists, a verification email has been sent."}
    if user.email_verified:
        return {"message": "Email already verified."}
    token = generate_token()
    user.verification_token = token
    db.commit()
    send_verification_email(email, token)
    return {"message": "Verification email sent. Check your inbox."}


@app.post("/api/forgot-password")
async def forgot_password(request: Request, db: Session = Depends(get_db)):
    from email_service import generate_token, send_password_reset_email
    body = await request.json()
    email = body.get("email", "").strip().lower()
    if not email:
        raise HTTPException(status_code=400, detail="Email required")

    # Rate limit: max 5 forgot-password requests per 15 minutes per email
    global _forgot_pw_last_cleanup
    now_ts = time_module.time()

    # Periodic cleanup: purge stale emails every 15 minutes
    if now_ts - _forgot_pw_last_cleanup > FORGOT_PW_RATE_WINDOW:
        stale = [k for k, v in _forgot_pw_timestamps.items() if not v or now_ts - v[-1] > FORGOT_PW_RATE_WINDOW]
        for k in stale:
            del _forgot_pw_timestamps[k]
        _forgot_pw_last_cleanup = now_ts

    timestamps = _forgot_pw_timestamps.get(email, [])
    timestamps = [t for t in timestamps if now_ts - t < FORGOT_PW_RATE_WINDOW]
    if len(timestamps) >= FORGOT_PW_RATE_LIMIT:
        return {"message": "If an account exists with that email, a password reset link has been sent."}
    timestamps.append(now_ts)
    _forgot_pw_timestamps[email] = timestamps

    user = db.query(User).filter(User.email == email).first()
    if user:
        token = generate_token()
        user.reset_token = token
        user.reset_token_expires = datetime.now(timezone.utc) + timedelta(hours=1)
        db.commit()
        send_password_reset_email(email, token)

    return {"message": "If an account exists with that email, a password reset link has been sent."}


@app.post("/api/reset-password")
async def reset_password(request: Request, db: Session = Depends(get_db)):
    body = await request.json()
    token = body.get("token", "")
    new_password = body.get("password", "")
    if not token or not new_password:
        raise HTTPException(status_code=400, detail="Token and new password required")
    if len(new_password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters")

    user = db.query(User).filter(User.reset_token == token).first()
    if not user:
        raise HTTPException(status_code=400, detail="Invalid or expired reset link")
    if user.reset_token_expires:
        expires = user.reset_token_expires if user.reset_token_expires.tzinfo else user.reset_token_expires.replace(tzinfo=timezone.utc)
        if expires < datetime.now(timezone.utc):
            raise HTTPException(status_code=400, detail="Reset link has expired. Request a new one.")

    user.password_hash = hash_password(new_password)
    user.reset_token = None
    user.reset_token_expires = None
    user.email_verified = True
    db.commit()
    return {"message": "Password reset successfully. You can now log in."}


@app.post("/api/login")
def login(req: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == req.email).first()
    if not user or not verify_password(req.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid credentials")

    # Require email verification (skip for admins and existing test accounts)
    if not getattr(user, 'email_verified', True) and user.role != "admin":
        raise HTTPException(status_code=403, detail="Please verify your email first. Check your inbox for the verification link.")

    token = create_access_token({
        "user_id": user.id,
        "role": user.role,
        "email": user.email
    })
    return {"access_token": token, "token_type": "bearer"}

# --- Profile Management ---
@app.get("/api/profile")
async def get_profile(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.id == user["user_id"]).first()
    if not db_user:
        raise HTTPException(status_code=404, detail="User not found")

    # Prefer base64 data (persistent) over file URL
    profile_pic = getattr(db_user, 'profile_picture_data', None)
    if not profile_pic:
        profile_pic = getattr(db_user, 'profile_picture', None)

    return {
        "email": db_user.email,
        "name": getattr(db_user, 'name', None),
        "studentId": getattr(db_user, 'student_id', None),
        "major": getattr(db_user, 'major', "Computer Science"),
        "profilePicture": profile_pic,
        "morganConnected": getattr(db_user, 'morgan_connected', False),
        "role": getattr(db_user, 'role', "student")
    }

@app.put("/api/profile")
async def update_profile(req: ProfileUpdateRequest, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.id == user["user_id"]).first()
    if not db_user: raise HTTPException(404, "User not found")
    
    if req.name is not None and hasattr(db_user, 'name'): db_user.name = req.name
    if req.studentId is not None and hasattr(db_user, 'student_id'): db_user.student_id = req.studentId
    if req.major is not None and hasattr(db_user, 'major'): db_user.major = req.major
    
    db.commit()
    return {"message": "Profile updated"}

@app.post("/api/change-password")
async def change_password(req: PasswordChangeRequest, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.id == user["user_id"]).first()
    if not db_user: raise HTTPException(404, "User not found")
    
    if not verify_password(req.currentPassword, db_user.password_hash):
        raise HTTPException(401, "Current password incorrect")
    
    db_user.password_hash = hash_password(req.newPassword)
    db.commit()
    return {"message": "Password changed"}

@app.post("/api/upload-profile-picture")
async def upload_profile_picture(profilePicture: UploadFile = File(...), user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    if not allowed_file(profilePicture.filename):
        raise HTTPException(400, "Invalid file type")

    # Read file content
    file_content = await profilePicture.read()

    # Get file extension and mime type
    ext = profilePicture.filename.rsplit('.', 1)[1].lower()
    mime_types = {
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'png': 'image/png',
        'gif': 'image/gif'
    }
    mime_type = mime_types.get(ext, 'image/jpeg')

    # Convert to base64 data URL
    import base64
    base64_data = base64.b64encode(file_content).decode('utf-8')
    data_url = f"data:{mime_type};base64,{base64_data}"

    # Also save to filesystem as backup
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"user_{user['user_id']}_{timestamp}.{ext}"
    filepath = os.path.join(UPLOAD_FOLDER, filename)

    with open(filepath, "wb") as f:
        f.write(file_content)

    file_url = f"/uploads/profile_pictures/{filename}"

    # Save base64 to database (persistent) and file URL as fallback
    db_user = db.query(User).filter(User.id == user["user_id"]).first()
    if db_user:
        db_user.profile_picture = file_url  # File path as fallback
        if hasattr(db_user, 'profile_picture_data'):
            db_user.profile_picture_data = data_url  # Base64 for persistence
        db.commit()

    # Return base64 data URL for immediate display
    return {"url": data_url}

#  NEW: Chat File Upload Endpoint
MAX_UPLOAD_SIZE = 10 * 1024 * 1024  # 10MB

@app.post("/api/upload-file")
async def upload_chat_file(file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    # 1. Validate File Type
    if not allowed_file(file.filename):
        raise HTTPException(400, "File type not allowed")

    # 2. Create Unique Filename
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    clean_name = re.sub(r'[^a-zA-Z0-9_.-]', '_', file.filename)
    filename = f"chat_{user['user_id']}_{timestamp}_{clean_name}"
    filepath = os.path.join(CHAT_FILES_FOLDER, filename)

    # 3. Stream to disk with size enforcement (never holds full file in memory)
    try:
        bytes_written = 0
        with open(filepath, "wb") as buffer:
            while chunk := await file.read(64 * 1024):  # 64KB chunks
                bytes_written += len(chunk)
                if bytes_written > MAX_UPLOAD_SIZE:
                    buffer.close()
                    os.remove(filepath)
                    raise HTTPException(413, f"File too large. Maximum size is {MAX_UPLOAD_SIZE // (1024*1024)}MB")
                buffer.write(chunk)
    except HTTPException:
        raise  # Preserve 413 for oversized files
    except Exception as e:
        print(f"[ERROR] File Save Error: {e}")
        raise HTTPException(500, "Could not save file")

    # 4. Return the public URL
    url = f"/uploads/chat_files/{filename}"
    return {"url": url, "filename": file.filename}

@app.post("/api/connect-morgan")
async def connect_morgan(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.id == user["user_id"]).first()
    if hasattr(db_user, 'morgan_connected'):
        db_user.morgan_connected = True
        db.commit()
    return {"message": "Morgan Connected", "morganConnected": True}

# ==============================================================================
# DegreeWorks Integration Endpoints
# ==============================================================================

@app.post("/api/degreeworks/sync")
async def sync_degreeworks(
    req: DegreeWorksRequest,
    user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Receives DegreeWorks data from the bookmarklet and saves it to the database.
    Creates or updates the user's DegreeWorks record.
    """
    try:
        db_user = db.query(User).filter(User.id == user["user_id"]).first()
        if not db_user:
            raise HTTPException(404, "User not found")

        # Check if user already has DegreeWorks data
        existing = db.query(DegreeWorksData).filter(DegreeWorksData.user_id == user["user_id"]).first()

        # Determine data source: only raw_data indicates a parsed/scraped source
        data_source = "manual_entry"
        if req.raw_data:
            data_source = "pdf_parse"

        if existing:
            # Update existing record
            existing.student_name = req.student_name
            existing.student_id = req.student_id
            existing.degree_program = req.degree_program
            existing.catalog_year = req.catalog_year
            existing.classification = req.classification
            existing.advisor = req.advisor
            existing.overall_gpa = req.overall_gpa
            existing.major_gpa = req.major_gpa
            existing.total_credits_earned = req.total_credits_earned
            existing.credits_required = req.credits_required
            existing.credits_remaining = req.credits_remaining
            existing.courses_completed = json.dumps(req.courses_completed) if req.courses_completed else None
            existing.courses_in_progress = json.dumps(req.courses_in_progress) if req.courses_in_progress else None
            existing.courses_remaining = json.dumps(req.courses_remaining) if req.courses_remaining else None
            existing.requirements_status = json.dumps(req.requirements_status) if req.requirements_status else None
            existing.raw_data = req.raw_data
            existing.data_source = data_source
            existing.updated_at = datetime.now(timezone.utc)
        else:
            # Create new record
            new_data = DegreeWorksData(
                user_id=user["user_id"],
                student_name=req.student_name,
                student_id=req.student_id,
                degree_program=req.degree_program,
                catalog_year=req.catalog_year,
                classification=req.classification,
                advisor=req.advisor,
                overall_gpa=req.overall_gpa,
                major_gpa=req.major_gpa,
                total_credits_earned=req.total_credits_earned,
                credits_required=req.credits_required,
                credits_remaining=req.credits_remaining,
                courses_completed=json.dumps(req.courses_completed) if req.courses_completed else None,
                courses_in_progress=json.dumps(req.courses_in_progress) if req.courses_in_progress else None,
                courses_remaining=json.dumps(req.courses_remaining) if req.courses_remaining else None,
                requirements_status=json.dumps(req.requirements_status) if req.requirements_status else None,
                raw_data=req.raw_data,
                data_source=data_source
            )
            db.add(new_data)

        # Update user's morgan_connected status
        db_user.morgan_connected = True
        db_user.morgan_connected_at = datetime.now(timezone.utc)

        # Update name if provided and not already set
        if req.student_name and not db_user.name:
            db_user.name = req.student_name
        if req.student_id and not db_user.student_id:
            db_user.student_id = req.student_id

        db.commit()

        return {
            "success": True,
            "message": "DegreeWorks data synced successfully!",
            "data": {
                "student_name": req.student_name,
                "degree_program": req.degree_program,
                "classification": req.classification,
                "gpa": req.overall_gpa,
                "credits_earned": req.total_credits_earned
            }
        }

    except Exception as e:
        print(f"[ERROR] DegreeWorks Sync Error: {e}")
        raise HTTPException(500, f"Failed to sync DegreeWorks data: {str(e)}")


@app.get("/api/degreeworks")
async def get_degreeworks(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """
    Retrieves the user's DegreeWorks data.
    """
    dw_data = db.query(DegreeWorksData).filter(DegreeWorksData.user_id == user["user_id"]).first()

    if not dw_data:
        return {"connected": False, "data": None}

    return {
        "connected": True,
        "data": {
            "student_name": dw_data.student_name,
            "student_id": dw_data.student_id,
            "degree_program": dw_data.degree_program,
            "catalog_year": dw_data.catalog_year,
            "classification": dw_data.classification,
            "advisor": dw_data.advisor,
            "overall_gpa": dw_data.overall_gpa,
            "major_gpa": dw_data.major_gpa,
            "total_credits_earned": dw_data.total_credits_earned,
            "credits_required": dw_data.credits_required,
            "credits_remaining": dw_data.credits_remaining,
            "courses_completed": json.loads(dw_data.courses_completed) if dw_data.courses_completed else [],
            "courses_in_progress": json.loads(dw_data.courses_in_progress) if dw_data.courses_in_progress else [],
            "courses_remaining": json.loads(dw_data.courses_remaining) if dw_data.courses_remaining else [],
            "requirements_status": json.loads(dw_data.requirements_status) if dw_data.requirements_status else [],
            "synced_at": dw_data.synced_at.isoformat() if dw_data.synced_at else None,
            "updated_at": dw_data.updated_at.isoformat() if dw_data.updated_at else None
        }
    }


@app.get("/api/degreeworks/debug")
async def debug_degreeworks(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """
    Debug endpoint to see ALL extracted DegreeWorks data including raw_data preview.
    """
    dw_data = db.query(DegreeWorksData).filter(DegreeWorksData.user_id == user["user_id"]).first()

    if not dw_data:
        return {"connected": False, "message": "No DegreeWorks data found for this user"}

    return {
        "connected": True,
        "all_fields": {
            "student_name": dw_data.student_name,
            "student_id": dw_data.student_id,
            "degree_program": dw_data.degree_program,
            "catalog_year": dw_data.catalog_year,
            "classification": dw_data.classification,
            "advisor": dw_data.advisor,
            "overall_gpa": dw_data.overall_gpa,
            "major_gpa": dw_data.major_gpa,
            "total_credits_earned": dw_data.total_credits_earned,
            "credits_required": dw_data.credits_required,
            "credits_remaining": dw_data.credits_remaining,
        },
        "courses_completed_count": len(json.loads(dw_data.courses_completed)) if dw_data.courses_completed else 0,
        "courses_completed": json.loads(dw_data.courses_completed) if dw_data.courses_completed else [],
        "raw_data_preview": dw_data.raw_data[:2000] if dw_data.raw_data else "No raw data",
        "raw_data_full": dw_data.raw_data[:10000] if dw_data.raw_data else "No raw data",
        "synced_at": dw_data.synced_at.isoformat() if dw_data.synced_at else None,
    }


@app.post("/api/degreeworks/test-pdf-parse")
async def test_pdf_parse(
    file: UploadFile = File(...),
    user: dict = Depends(get_current_user)
):
    """
    Test endpoint that parses a DegreeWorks PDF and returns what was extracted
    WITHOUT saving to database. Useful for debugging.
    """
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(400, "Please upload a PDF file")

    try:
        # Save temporarily
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        temp_filename = f"test_dw_{user['user_id']}_{timestamp}.pdf"
        temp_filepath = os.path.join(CHAT_FILES_FOLDER, temp_filename)

        with open(temp_filepath, "wb") as buffer:
            content = await file.read()
            buffer.write(content)

        # Extract text from PDF
        pdf_text = ""
        try:
            reader = pypdf.PdfReader(temp_filepath)
            for page in reader.pages:
                pdf_text += page.extract_text() + "\n"
        except Exception as e:
            return {"error": f"Could not read PDF: {e}"}

        # Parse the PDF
        data = parse_degreeworks_pdf(pdf_text)

        # Clean up temp file
        try:
            os.remove(temp_filepath)
        except:
            pass

        return {
            "success": True,
            "pdf_text_length": len(pdf_text),
            "pdf_text_preview": pdf_text[:3000],
            "extracted_data": {
                "student_name": data.get('student_name'),
                "student_id": data.get('student_id'),
                "classification": data.get('classification'),
                "degree_program": data.get('degree_program'),
                "overall_gpa": data.get('overall_gpa'),
                "major_gpa": data.get('major_gpa'),
                "total_credits_earned": data.get('total_credits_earned'),
                "credits_required": data.get('credits_required'),
                "credits_remaining": data.get('credits_remaining'),
                "advisor": data.get('advisor'),
                "catalog_year": data.get('catalog_year'),
                "courses_count": len(json.loads(data.get('courses_completed', '[]'))) if data.get('courses_completed') else 0
            },
            "message": "Test parse complete - data NOT saved to database"
        }

    except Exception as e:
        return {"error": f"Failed to process PDF: {str(e)}"}


@app.delete("/api/degreeworks/disconnect")
async def disconnect_degreeworks(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """
    Removes the user's DegreeWorks data and disconnects their Morgan account.
    """
    try:
        # Delete DegreeWorks data
        db.query(DegreeWorksData).filter(DegreeWorksData.user_id == user["user_id"]).delete()

        # Update user's morgan_connected status
        db_user = db.query(User).filter(User.id == user["user_id"]).first()
        if db_user:
            db_user.morgan_connected = False
            db_user.morgan_connected_at = None

        db.commit()

        return {"success": True, "message": "DegreeWorks data disconnected"}
    except Exception as e:
        print(f"[ERROR] DegreeWorks Disconnect Error: {e}")
        raise HTTPException(500, f"Failed to disconnect: {str(e)}")


@app.post("/api/degreeworks/upload-pdf")
async def upload_degreeworks_pdf(
    file: UploadFile = File(...),
    user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Uploads DegreeWorks document (PDF or DOCX) and stores the extracted
    text for chat context injection.
    """
    ALLOWED_DW_EXTENSIONS = {'pdf', 'docx', 'doc'}

    print("=" * 60)
    print("DEGREEWORKS UPLOAD ENDPOINT HIT!")
    print(f"File received: {file.filename if file else 'NO FILE'}")
    print(f"User: {user}")
    print("=" * 60)

    if not file or not file.filename:
        raise HTTPException(400, "No file provided")

    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in ALLOWED_DW_EXTENSIONS:
        raise HTTPException(400, f"Unsupported file type. Please upload: {', '.join(ALLOWED_DW_EXTENSIONS)}")

    try:
        # Save the uploaded file temporarily
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        temp_filename = f"degreeworks_{user['user_id']}_{timestamp}.{ext}"
        temp_filepath = os.path.join(CHAT_FILES_FOLDER, temp_filename)

        content = await file.read()
        print(f"Received file: {file.filename}, size: {len(content)} bytes")

        with open(temp_filepath, "wb") as buffer:
            buffer.write(content)

        # Extract text - try fast local methods first, OCR API only when needed
        pdf_text = ""

        # Method 1: Local pypdf for PDFs (instant for text-based PDFs)
        if ext == 'pdf':
            try:
                print("Trying local pypdf extraction (fast)...")
                reader = pypdf.PdfReader(temp_filepath)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pdf_text += page_text + "\n"
                print(f"pypdf extracted {len(pdf_text)} chars")
            except Exception as e:
                print(f"pypdf extraction failed: {e}")

        # Method 2: Local python-docx for DOCX (instant)
        if ext in ('docx', 'doc'):
            try:
                print("Trying local docx extraction (fast)...")
                doc_file = docx.Document(temp_filepath)
                for para in doc_file.paragraphs:
                    pdf_text += para.text + "\n"
                print(f"docx extracted {len(pdf_text)} chars")
            except Exception as e:
                print(f"docx extraction failed: {e}")

        print(f"Total extracted text: {len(pdf_text)} characters")

        if len(pdf_text.strip()) < 20:
            raise HTTPException(
                400,
                f"Could not extract text from this file ({len(pdf_text)} chars). "
                "Please upload a text-based PDF or DOCX file."
            )

        # Try to parse specific fields (best effort)
        data = parse_degreeworks_pdf(pdf_text)

        # CRITICAL: Always store the raw PDF text - this is used for chat context injection
        data['raw_data'] = pdf_text[:50000]  # Store up to 50k chars

        # Get or create DegreeWorks record
        db_user = db.query(User).filter(User.id == user["user_id"]).first()
        existing = db.query(DegreeWorksData).filter(DegreeWorksData.user_id == user["user_id"]).first()

        if existing:
            # Update existing - ALWAYS update raw_data
            existing.raw_data = data['raw_data']
            for key, value in data.items():
                if value is not None and hasattr(existing, key):
                    setattr(existing, key, value)
            existing.updated_at = datetime.now(timezone.utc)
        else:
            # Create new
            new_data = DegreeWorksData(user_id=user["user_id"], **data)
            db.add(new_data)

        # Update user's morgan_connected status
        db_user.morgan_connected = True
        db_user.morgan_connected_at = datetime.now(timezone.utc)

        # Update user name if found
        if data.get('student_name') and not db_user.name:
            db_user.name = data['student_name']

        db.commit()

        # Clean up temp file
        try:
            os.remove(temp_filepath)
        except:
            pass

        return {
            "success": True,
            "message": "DegreeWorks PDF uploaded successfully! Your academic data is now available for personalized chat.",
            "data": {
                "student_name": data.get('student_name'),
                "classification": data.get('classification'),
                "degree_program": data.get('degree_program'),
                "overall_gpa": data.get('overall_gpa'),
                "total_credits_earned": data.get('total_credits_earned'),
                "pdf_text_length": len(pdf_text),
                "pdf_stored": True
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] DegreeWorks PDF Upload Error: {e}")
        raise HTTPException(500, f"Failed to process PDF: {str(e)}")


@app.get("/api/bookmarklet.js")
async def serve_bookmarklet_js(token: str = "", api: str = ""):
    """
    Serves the full bookmarklet sync script.
    The bookmarklet itself is a tiny loader that injects this script.
    This avoids Chrome's ~2KB URL length limit for bookmarklets.
    """
    # Escape values to prevent JS injection via crafted query params
    safe_api = json.dumps(api)    # wraps in quotes + escapes special chars
    safe_token = json.dumps(token)
    js_code = f"""
(function() {{
  var API = {safe_api};
  var TOKEN = {safe_token};

  var msg = document.createElement('div');
  msg.style.cssText = 'position:fixed;top:20px;right:20px;background:#002D72;color:#fff;padding:20px 24px;border-radius:12px;z-index:999999;font-family:system-ui,Arial;font-size:14px;box-shadow:0 4px 20px rgba(0,0,0,0.3);';
  msg.innerHTML = '<strong>CS Navigator</strong><br>Syncing your DegreeWorks data...';
  document.body.appendChild(msg);

  var html = document.documentElement.outerHTML;

  fetch(API + '/api/degreeworks/scrape-html', {{
    method: 'POST',
    headers: {{
      'Content-Type': 'application/json',
      'Authorization': 'Bearer ' + TOKEN
    }},
    body: JSON.stringify({{ html: html }})
  }})
  .then(function(r) {{ return r.json(); }})
  .then(function(d) {{
    msg.remove();
    if (d.success) {{
      var info = d.data || {{}};
      var details = '';
      if (info.overall_gpa) details += 'GPA: ' + info.overall_gpa + '\\n';
      if (info.classification) details += 'Classification: ' + info.classification + '\\n';
      if (info.total_credits_earned) details += 'Credits: ' + info.total_credits_earned + '\\n';
      if (info.courses_count) details += 'Courses found: ' + info.courses_count + '\\n';
      alert('DegreeWorks synced successfully!\\n\\n' + details + '\\nYou can now close this tab and return to CS Navigator.');
    }} else {{
      alert('Sync failed: ' + (d.detail || d.message || 'Unknown error') + '\\n\\nTry using manual entry instead.');
    }}
  }})
  .catch(function(e) {{
    msg.remove();
    alert('Error: ' + e.message + '\\n\\nTry using manual entry instead.');
  }});
}})();
"""
    from fastapi.responses import Response
    return Response(content=js_code, media_type="application/javascript")


@app.post("/api/degreeworks/scrape-html")
async def scrape_degreeworks_html(
    request: dict,
    user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Receives raw HTML from DegreeWorks page (via bookmarklet) and extracts academic data.
    This is the preferred method - more accurate than PDF parsing.
    """
    html_content = request.get("html", "")

    if not html_content or len(html_content) < 100:
        raise HTTPException(400, "No HTML content provided")

    try:
        # Parse the HTML and extract data
        data = parse_degreeworks_html(html_content)

        if not data or (not data.get('overall_gpa') and not data.get('classification')):
            raise HTTPException(400, "Could not extract academic data from this page. Make sure you're on the DegreeWorks audit page.")

        # Get or create DegreeWorks record
        db_user = db.query(User).filter(User.id == user["user_id"]).first()
        existing = db.query(DegreeWorksData).filter(DegreeWorksData.user_id == user["user_id"]).first()

        if existing:
            # Update existing
            for key, value in data.items():
                if value is not None and hasattr(existing, key):
                    setattr(existing, key, value)
            existing.updated_at = datetime.now(timezone.utc)
        else:
            # Create new
            new_data = DegreeWorksData(user_id=user["user_id"], **data)
            db.add(new_data)

        # Update user's morgan_connected status
        db_user.morgan_connected = True
        db_user.morgan_connected_at = datetime.now(timezone.utc)

        # Update user name if found
        if data.get('student_name') and not db_user.name:
            db_user.name = data['student_name']
        if data.get('student_id') and not db_user.student_id:
            db_user.student_id = data['student_id']

        db.commit()

        return {
            "success": True,
            "message": "DegreeWorks data synced successfully!",
            "data": {
                "student_name": data.get('student_name'),
                "classification": data.get('classification'),
                "degree_program": data.get('degree_program'),
                "overall_gpa": data.get('overall_gpa'),
                "major_gpa": data.get('major_gpa'),
                "total_credits_earned": data.get('total_credits_earned'),
                "courses_count": len(json.loads(data.get('courses_completed', '[]'))) if data.get('courses_completed') else 0
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] DegreeWorks HTML Scrape Error: {e}")
        raise HTTPException(500, f"Failed to process DegreeWorks data: {str(e)}")


def parse_degreeworks_html(html: str) -> dict:
    """
    Parses DegreeWorks HTML and extracts academic data.
    Based on the DegreeWorks HTML structure.
    """
    data = {}

    # First, extract all text content from HTML
    # Remove script and style tags
    html_clean = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
    html_clean = re.sub(r'<style[^>]*>.*?</style>', '', html_clean, flags=re.DOTALL | re.IGNORECASE)
    # Remove HTML tags but keep text
    text = re.sub(r'<[^>]+>', ' ', html_clean)
    # Clean up whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    text_lower = text.lower()

    # Also keep original HTML for pattern matching
    html_lower = html.lower()

    # Extract Student Name - look in multiple places
    # Blacklist common false positives from DegreeWorks labels
    name_blacklist = {'academic standing', 'good standing', 'computer science', 'general education',
                      'bachelor of science', 'student self service', 'morgan state', 'degree works',
                      'natural science', 'free electives', 'university requirements'}
    name_patterns = [
        r'student[:\s]+name[:\s]+([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?(?:\s+[A-Z][a-z]+)+)',
        # DegreeWorks profile area: name appears after "Profile" aria label
        r'Profile["\s>]+[^<]*>([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?(?:\s+[A-Z][a-z]+)+)<',
        # Hello greeting from Banner
        r'Hello\s+([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?(?:\s+[A-Z][a-z]+)+)',
        r'>([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?(?:\s+[A-Z][a-z]+)+)</(?:h1|h2|div|span)',
        r'name[:\s]+([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?(?:\s+[A-Z][a-z]+)+)',
    ]
    for pattern in name_patterns:
        match = re.search(pattern, html)
        if match:
            name = match.group(1).strip()
            if len(name) > 3 and len(name) < 50 and name.lower() not in name_blacklist:
                data['student_name'] = name
                break

    # Student ID
    id_patterns = [
        r'(?:student\s*)?id[:\s#]*["\']?(\d{7,9})["\']?',
        r'>\s*(\d{7,9})\s*<',
    ]
    for pattern in id_patterns:
        match = re.search(pattern, html_lower)
        if match:
            data['student_id'] = match.group(1)
            break

    # Overall GPA - this is critical (support 1-3 decimal places)
    gpa_patterns = [
        r'overall\s*gpa[:\s]*(\d\.\d{1,3})',
        r'cumulative\s*gpa[:\s]*(\d\.\d{1,3})',
        r'gpa[:\s]+(\d\.\d{1,3})',
        r'gpa\s*:\s*(\d\.\d{1,3})',
        r'gpa\s+(\d\.\d{1,3})',
        r'total\s*gpa[:\s]*(\d\.\d{1,3})',
        r'career\s*gpa[:\s]*(\d\.\d{1,3})',
        r'>(\d\.\d{2,3})<.*?gpa',
        r'gpa.*?>(\d\.\d{2,3})<',
    ]
    for pattern in gpa_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                gpa = float(match.group(1))
                if 0 <= gpa <= 4.0:
                    data['overall_gpa'] = gpa
                    break
            except:
                pass

    # Fallback: Find any decimal that looks like GPA
    if not data.get('overall_gpa'):
        # Look near GPA keyword
        gpa_area = re.search(r'gpa.{0,30}', text_lower)
        if gpa_area:
            gpa_match = re.search(r'(\d\.\d{1,3})', gpa_area.group())
            if gpa_match:
                gpa = float(gpa_match.group(1))
                if 1.0 <= gpa <= 4.0:
                    data['overall_gpa'] = gpa

    # Major GPA (support 1-3 decimal places)
    major_gpa_patterns = [
        r'major\s*gpa[:\s]*(\d\.\d{1,3})',
        r'program\s*gpa[:\s]*(\d\.\d{1,3})',
    ]
    for pattern in major_gpa_patterns:
        match = re.search(pattern, text_lower)
        if match:
            data['major_gpa'] = float(match.group(1))
            break

    # Classification - handle DegreeWorks format like "4-Senior" or bare "Senior"
    class_patterns = [
        r'classification[:\s]*(?:\d[- ])?(freshman|sophomore|junior|senior|graduate)',
        r'class[:\s]*(?:\d[- ])?(freshman|sophomore|junior|senior|graduate)',
        r'standing[:\s]*(?:\d[- ])?(freshman|sophomore|junior|senior|graduate)',
        r'level[:\s]*(?:\d[- ])?(freshman|sophomore|junior|senior|graduate)',
        r'student\s*level[:\s]*(?:\d[- ])?(freshman|sophomore|junior|senior|graduate)',
        r'>(?:\d[- ])?(freshman|sophomore|junior|senior|graduate)<',
    ]
    for pattern in class_patterns:
        match = re.search(pattern, text_lower)
        if match:
            data['classification'] = match.group(1).title()
            break

    # If classification not found, try to determine from credits
    if not data.get('classification') and data.get('total_credits_earned'):
        credits = data['total_credits_earned']
        if credits >= 90:
            data['classification'] = 'Senior'
        elif credits >= 60:
            data['classification'] = 'Junior'
        elif credits >= 30:
            data['classification'] = 'Sophomore'
        else:
            data['classification'] = 'Freshman'

    # Degree/Program/Major
    degree_patterns = [
        r'(bachelor\s+of\s+science\s+in\s+[a-z\s]+?)(?:\s|<|$)',
        r'(master\s+of\s+science\s+in\s+[a-z\s]+?)(?:\s|<|$)',
        r'major[:\s]*(computer\s+science|information\s+(?:systems|science)|cybersecurity|software\s+engineering)',
        r'program[:\s]*(computer\s+science|information\s+(?:systems|science)|cybersecurity)',
        r'>(computer\s+science|information\s+systems)<',
    ]
    for pattern in degree_patterns:
        match = re.search(pattern, text_lower)
        if match:
            program = match.group(1).strip().title()
            if len(program) > 5:
                data['degree_program'] = program
                break

    # Credits - look for various patterns (DegreeWorks uses "Credits applied")
    credits_patterns = [
        r'credits\s*applied[:\s]*(\d{2,3}(?:\.\d)?)',
        r'(?:total|earned|completed)\s*(?:credits|hours)[:\s]*(\d{2,3}(?:\.\d)?)',
        r'(\d{2,3}(?:\.\d)?)\s*(?:credits|hours)\s*(?:earned|completed|applied|total)',
        r'credits\s*(?:earned|applied)[:\s]*(\d{2,3}(?:\.\d)?)',
        r'hours\s*earned[:\s]*(\d{2,3}(?:\.\d)?)',
        r'credits\s*:\s*(\d{2,3}(?:\.\d)?)',
        r'total\s*credits[:\s]*(\d{2,3}(?:\.\d)?)',
    ]
    for pattern in credits_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                credits = float(match.group(1))
                # Valid range for student credits (20-200), avoid 100 (often percentage)
                if 20 < credits <= 200 and credits != 100:
                    data['total_credits_earned'] = credits
                    break
            except:
                pass

    # Credits Required
    req_patterns = [
        r'(?:credits|hours)\s*required[:\s]*(\d{2,3})',
        r'required[:\s]*(\d{2,3})\s*(?:credits|hours)',
        r'total\s*(?:credits|hours)[:\s]*(\d{2,3})',
    ]
    for pattern in req_patterns:
        match = re.search(pattern, text_lower)
        if match:
            req = float(match.group(1))
            if 60 <= req <= 180:
                data['credits_required'] = req
                break

    # Credits Remaining
    remain_patterns = [
        r'(?:remaining|still\s*need|left)[:\s]*(\d{1,3}(?:\.\d)?)\s*(?:credits|hours)?',
        r'(\d{1,3}(?:\.\d)?)\s*(?:credits|hours)\s*(?:remaining|left|needed)',
    ]
    for pattern in remain_patterns:
        match = re.search(pattern, text_lower)
        if match:
            remain = float(match.group(1))
            if 0 <= remain <= 150:
                data['credits_remaining'] = remain
                break

    # Catalog Year - handle "SPRING 2024" or "2024-2025" or just "2024"
    catalog_match = re.search(r'catalog\s*year[:\s]*((?:spring|summer|fall|winter)\s+\d{4})', text_lower)
    if catalog_match:
        data['catalog_year'] = catalog_match.group(1).title()
    else:
        catalog_match = re.search(r'(?:catalog|requirement)[:\s]*(\d{4}[-\u2013]\d{4}|\d{4})', text_lower)
        if catalog_match:
            data['catalog_year'] = catalog_match.group(1)

    # Advisor - require at least 2 words (first + last name), blacklist common false positives
    advisor_blacklist = {'prerequisite', 'self service', 'good standing', 'academic standing',
                         'computer science', 'general education'}
    advisor_patterns = [
        r'advisor[:\s]+([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?(?:\s+[A-Z][a-z]+)+)',
        r'advised\s+by[:\s]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
    ]
    for pattern in advisor_patterns:
        match = re.search(pattern, text)
        if match:
            advisor = match.group(1).strip()
            if len(advisor) > 3 and advisor.lower() not in advisor_blacklist:
                data['advisor'] = advisor
                break

    # Extract completed courses - look for course patterns
    # Pattern: DEPT 123 Course Title Grade Credits Term
    # DegreeWorks uses UPPERCASE terms (SPRING 2024) and TRA for transfer grades
    VALID_GRADES = ['A', 'A-', 'A+', 'B', 'B-', 'B+', 'C', 'C-', 'C+', 'D', 'D-', 'D+', 'F',
                     'TRA', 'TRB', 'TRC', 'TRD', 'TR', 'P', 'S', 'IP', 'W']
    courses_completed = []
    course_patterns = [
        # Standard format: COSC 111 INTRO TO COMPUTER SCI I A 3 SPRING 2024
        # Course numbers can have up to 2 letter suffixes (e.g. PHYS 116TR)
        r'([A-Z]{2,4})\s*(\d{3}[A-Z]{0,2})\s+([A-Za-z][A-Za-z\s&\-,\(\)\']+?)\s+([ABCDF][+-]?|TR[A-D]|TR|IP|P|S|W)\s+(\d{1,2}(?:\.\d{1,2})?)\s+((?:FALL|SPRING|SUMMER|WINTER|Fall|Spring|Summer|Winter)\s*\d{4})',
        # Without term: COSC 111 INTRO TO CS A 3
        r'([A-Z]{2,4})\s*(\d{3}[A-Z]{0,2})\s+([A-Za-z][A-Za-z\s&\-,\(\)\']{3,40}?)\s+([ABCDF][+-]?|TR[A-D]|TR|IP|P|S|W)\s+(\d{1,2}(?:\.\d{1,2})?)',
        # Minimal: COSC 111 A 3
        r'([A-Z]{2,4})\s*(\d{3}[A-Z]{0,2})\s+([ABCDF][+-]?|TR[A-D]|TR|IP|P|S|W)\s+(\d{1,2}(?:\.\d{1,2})?)',
    ]

    seen_courses = set()
    for pattern in course_patterns:
        for match in re.finditer(pattern, text):
            groups = match.groups()
            if len(groups) >= 4:
                dept = groups[0]
                num = groups[1]
                code = f"{dept} {num}"

                if code in seen_courses:
                    continue
                seen_courses.add(code)

                course = {'code': code}

                if len(groups) == 6:  # Full format with term
                    course['name'] = groups[2].strip()[:50]
                    course['grade'] = groups[3]
                    course['credits'] = float(groups[4])
                    course['term'] = groups[5].title()
                elif len(groups) == 5:  # Without term
                    course['name'] = groups[2].strip()[:50]
                    course['grade'] = groups[3]
                    course['credits'] = float(groups[4])
                elif len(groups) == 4:  # Minimal
                    course['grade'] = groups[2]
                    course['credits'] = float(groups[3])

                if course.get('grade') in VALID_GRADES:
                    courses_completed.append(course)

    if courses_completed:
        data['courses_completed'] = json.dumps(courses_completed[:50])  # Limit to 50 courses

    # Extract remaining/needed courses
    remaining_courses = []
    still_needed_pattern = r'still\s+need(?:ed)?[:\s]+(.+?)(?:\.|$)'
    for match in re.finditer(still_needed_pattern, text_lower):
        req = match.group(1).strip()
        if len(req) > 5 and len(req) < 200:
            remaining_courses.append({'requirement': req})

    if remaining_courses:
        data['courses_remaining'] = json.dumps(remaining_courses[:20])

    # Store raw text for reference (truncated)
    data['raw_data'] = text[:30000]

    return data


def parse_degreeworks_pdf(text: str) -> dict:
    """
    Parses DegreeWorks PDF text using pure text processing.
    No LLM needed - cleans the text first, then extracts structured data with regex.
    Fast, deterministic, and free (no API call).
    """
    data = {}

    # Store raw text for the "chat with PDF" feature
    data['raw_data'] = text[:30000]

    # =====================================================
    # STEP 1: Clean the raw PDF text
    # Remove noise, collapse multi-line entries, keep only useful lines
    # =====================================================
    lines = text.split('\n')
    clean_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Skip noise lines
        if stripped.startswith('Satisfied by:'):
            continue
        if stripped.startswith('Exception by:'):
            continue
        if stripped.startswith('Morgan State University') and '- *****' in stripped:
            continue
        if stripped.startswith('Disclaimer'):
            break
        if stripped.startswith('Legend'):
            break
        if stripped.startswith('Ellucian Degree'):
            break
        clean_lines.append(stripped)

    clean_text = '\n'.join(clean_lines)
    # Also make a single-line version for multi-line course matching
    collapsed = ' '.join(clean_lines)

    print("=" * 60)
    print(f"PDF: {len(text)} chars raw -> {len(clean_text)} chars cleaned")
    print("=" * 60)

    # =====================================================
    # STEP 2: Extract header fields (GPA, name, classification, etc.)
    # =====================================================

    # Student name: "Student name Last, First"
    name_match = re.search(r'Student\s+name\s+(\w[\w\'-]+),\s+(\w[\w\'-]+)', text)
    if name_match:
        data['student_name'] = f"{name_match.group(2)} {name_match.group(1)}"

    # Overall GPA: "Overall GPA\n3.953" or "GPA: 3.953"
    gpa_match = re.search(r'Overall\s+GPA\s*[:\n]?\s*(\d\.\d{1,3})', text)
    if gpa_match:
        gpa = float(gpa_match.group(1))
        if 0.0 <= gpa <= 4.0:
            data['overall_gpa'] = gpa

    # Major GPA: "Your GPA in these classes is 4.000"
    major_gpa_match = re.search(r'Your\s+GPA\s+in\s+these\s+classes\s+is\s+(\d\.\d{1,3})', text)
    if major_gpa_match:
        mgpa = float(major_gpa_match.group(1))
        if 0.0 <= mgpa <= 4.0:
            data['major_gpa'] = mgpa

    # Classification: "Classification 4-Senior" or "Classification Senior"
    class_match = re.search(r'Classification\s+(?:\d-)?(Freshman|Sophomore|Junior|Senior|Graduate)', text, re.IGNORECASE)
    if class_match:
        data['classification'] = class_match.group(1).title()

    # Credits applied: "Credits applied:  128.5"
    credits_match = re.search(r'Credits\s+applied:\s*(\d+\.?\d*)', text)
    if credits_match:
        creds = float(credits_match.group(1))
        if 0 <= creds <= 300:
            data['total_credits_earned'] = creds

    # Credits required: "Credits required: 120"
    creq_match = re.search(r'Credits\s+required:\s*(\d+\.?\d*)', text)
    if creq_match:
        creq = float(creq_match.group(1))
        if 30 <= creq <= 300:
            data['credits_required'] = creq
            if data.get('total_credits_earned'):
                remaining = max(0, creq - data['total_credits_earned'])
                data['credits_remaining'] = remaining

    # Degree program: "Degree Bachelor of Science" + "Major Computer Science"
    degree_match = re.search(r'Degree\s+(Bachelor\s+of\s+\w+|Master\s+of\s+\w+)', text)
    major_match = re.search(r'Major\s+([A-Za-z ]+?)(?:\s{2,}|Program)', text)
    if degree_match and major_match:
        data['degree_program'] = f"{degree_match.group(1)} in {major_match.group(1).strip()}"
    elif degree_match:
        data['degree_program'] = degree_match.group(1)

    # Advisor: "Advisor Vojislav Stojkovic" (stop at double-space or end of line)
    advisor_match = re.search(r'Advisor\s+([A-Z][a-z]+(?:\s[A-Z][a-z]+)+)', text)
    if advisor_match:
        data['advisor'] = advisor_match.group(1).strip()

    # Catalog year: "Catalog year:  SPRING 2024"
    catalog_match = re.search(r'Catalog\s+year:\s*(\w+\s+\d{4})', text)
    if catalog_match:
        data['catalog_year'] = catalog_match.group(1)

    # Transfer hours (extracted but not stored in DB - kept in raw_data only)
    # transfer_match = re.search(r'Transfer\s*Hours\s+(\d+\.?\d*)', text)

    # =====================================================
    # STEP 3: Extract ALL courses from cleaned collapsed text
    # Pattern: DEPT CODE  COURSE NAME  GRADE  CREDITS  TERM
    # Handles multi-line names because text is collapsed
    # =====================================================

    # Course code prefixes we care about (add more as needed)
    DEPT_PREFIXES = r'(?:COSC|MATH|CLCO|EEGR|INSS|PHYS|BIOL|CHEM|ENGL|HIST|PSYC|PHIL|HLTH|WGST|FIN|ORTR|THEA|PHEC)'

    # Letter grades, transfer grades, pass/fail, and in-progress
    VALID_GRADES = r'(?:A\+?|A-|B\+?|B-|C\+?|C-|D\+?|D-|F|TRA|TRB|TRC|TRD|PT|IP|W)'

    # Main course extraction pattern on collapsed text
    # Course name: up to ~60 chars of letters/digits/spaces/punctuation, but NOT containing
    # another course code or grade-like pattern (prevents runaway matching)
    course_pattern = re.compile(
        r'(' + DEPT_PREFIXES + r'\s+\d{3}(?:TR)?)\s+'  # course code (e.g., COSC 470, PHYS 116TR)
        r'([A-Z][A-Za-z0-9 &/\',\.\-\(\)]{2,55}?)\s+'  # course name (2-55 chars, starts with uppercase)
        r'\b(' + VALID_GRADES + r')\b\s+'                 # grade with word boundary
        r'(\d+\.?\d*)\s+'                                  # credits
        r'((?:FALL|SPRING|SUMMER)\s+\d{4})',              # term
        re.IGNORECASE
    )

    # In-progress pattern: "COSC 458 SOFTWARE ENGINEERING IP (3) SPRING 2026"
    # Course name limited to 55 chars max to prevent runaway across multiple entries
    ip_pattern = re.compile(
        r'(' + DEPT_PREFIXES + r'\s+\d{3})\s+'
        r'([A-Z][A-Za-z0-9 &/\',\.\-\(\)]{2,55}?)\s+'
        r'IP\s+\((\d+)\)\s+'
        r'((?:FALL|SPRING|SUMMER)\s+\d{4})',
        re.IGNORECASE
    )

    completed_courses = []
    ip_courses = []
    seen_codes = set()

    # First pass: extract in-progress courses (IP pattern is more specific)
    for match in ip_pattern.finditer(collapsed):
        code = match.group(1).upper().strip()
        name = match.group(2).strip()
        credits = int(match.group(3))
        term = match.group(4).strip()
        if code not in seen_codes:
            seen_codes.add(code)
            ip_courses.append({
                "code": code,
                "name": name,
                "credits": credits,
                "status": "in_progress",
                "term": term
            })

    # Second pass: extract completed courses
    for match in course_pattern.finditer(collapsed):
        code = match.group(1).upper().strip()
        name = match.group(2).strip()
        grade = match.group(3).upper().strip()
        credits = float(match.group(4))
        term = match.group(5).strip()

        if code in seen_codes:
            continue
        seen_codes.add(code)

        if grade == 'IP':
            ip_courses.append({
                "code": code,
                "name": name,
                "credits": int(credits),
                "status": "in_progress",
                "term": term
            })
        else:
            completed_courses.append({
                "code": code,
                "name": name,
                "grade": grade,
                "credits": credits,
                "term": term
            })

    if completed_courses:
        data['courses_completed'] = json.dumps(completed_courses)
    if ip_courses:
        data['courses_in_progress'] = json.dumps(ip_courses)

    # Derive classification from credits if not found in header
    if not data.get("classification") and data.get("total_credits_earned"):
        credits = data["total_credits_earned"]
        if credits >= 90:
            data["classification"] = "Senior"
        elif credits >= 60:
            data["classification"] = "Junior"
        elif credits >= 30:
            data["classification"] = "Sophomore"
        else:
            data["classification"] = "Freshman"

    print("=" * 60)
    print("EXTRACTION SUMMARY:")
    print(f"   Name: {data.get('student_name', 'NOT FOUND')}")
    print(f"   GPA: {data.get('overall_gpa', 'NOT FOUND')}")
    print(f"   Major GPA: {data.get('major_gpa', 'NOT FOUND')}")
    print(f"   Credits: {data.get('total_credits_earned', 'NOT FOUND')}")
    print(f"   Classification: {data.get('classification', 'NOT FOUND')}")
    print(f"   Program: {data.get('degree_program', 'NOT FOUND')}")
    print(f"   Advisor: {data.get('advisor', 'NOT FOUND')}")
    print(f"   Courses Completed: {len(completed_courses)}")
    print(f"   Courses In Progress: {len(ip_courses)}")
    if completed_courses:
        print(f"   Completed codes: {[c['code'] for c in completed_courses]}")
    if ip_courses:
        print(f"   In-progress codes: {[c['code'] for c in ip_courses]}")
    print("=" * 60)

    return data


# ==============================================================================
# Banner Student Self Service Integration Endpoints
# ==============================================================================

class BannerSyncRequest(BaseModel):
    """Request body for Banner SSB sync. Credentials are in-memory only."""
    username: str
    password: str


@app.post("/api/banner/sync")
async def sync_banner_data(
    req: BannerSyncRequest,
    user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Full Banner SSB sync via CAS authentication.
    Authenticates with MSU CAS, calls Banner REST APIs,
    updates DegreeWorksData + BannerStudentData in RDS.
    Returns SSE progress stream.
    """
    user_id = user["user_id"]

    # Rate limit: max 3 syncs per user per hour
    now = datetime.now(timezone.utc)
    timestamps = _banner_sync_timestamps.get(user_id, [])
    one_hour_ago = now.timestamp() - 3600
    timestamps = [t for t in timestamps if t > one_hour_ago]
    if len(timestamps) >= 3:
        raise HTTPException(429, "Rate limit exceeded. Maximum 3 syncs per hour.")
    timestamps.append(now.timestamp())
    _banner_sync_timestamps[user_id] = timestamps

    async def generate_sse():
        """SSE stream for sync progress."""
        try:
            progress_steps = []

            async def track_progress(step, detail):
                progress_steps.append({"step": step, "detail": detail})

            # Run the sync (DegreeWorks + Student Profile)
            results = await sync_banner(req.username, req.password, track_progress)

            # Stream progress steps
            for p in progress_steps:
                yield f"data: {json.dumps({'type': 'progress', 'step': p['step'], 'detail': p['detail']})}\n\n"

            # Process results and update database
            sync_db = SessionLocal()
            try:
                db_user = sync_db.query(User).filter(User.id == user_id).first()
                if not db_user:
                    yield f"data: {json.dumps({'type': 'error', 'detail': 'User not found'})}\n\n"
                    return

                yield f"data: {json.dumps({'type': 'progress', 'step': 'saving', 'detail': 'Saving to database...'})}\n\n"

                existing_dw = sync_db.query(DegreeWorksData).filter(DegreeWorksData.user_id == user_id).first()
                if not existing_dw:
                    existing_dw = DegreeWorksData(user_id=user_id)
                    sync_db.add(existing_dw)

                # 1. Parse DegreeWorks JSON audit (primary, richest source)
                dw_json = results.get("degreeworks_json")
                dw_data = {}
                if dw_json:
                    try:
                        from banner_scraper.parsers import parse_degreeworks_audit_json
                        dw_data = parse_degreeworks_audit_json(dw_json)
                        # Apply all DW fields
                        for key, value in dw_data.items():
                            if value is not None and hasattr(existing_dw, key):
                                setattr(existing_dw, key, value)
                    except Exception as e:
                        print(f"[BANNER] DW JSON parse error: {e}")

                # 2. Parse Student Profile HTML (fills gaps DW might miss)
                profile_html = results.get("profile_html")
                profile = {}
                if profile_html:
                    try:
                        from banner_scraper.parsers import parse_student_profile
                        profile = parse_student_profile({"type": "html", "data": profile_html})
                        # Only fill in gaps (DW data takes priority)
                        if not existing_dw.student_name and profile.get("name"):
                            existing_dw.student_name = profile["name"]
                        if not existing_dw.student_id and profile.get("student_id"):
                            existing_dw.student_id = profile["student_id"]
                        if not existing_dw.classification and profile.get("classification"):
                            existing_dw.classification = profile["classification"]
                        if not existing_dw.advisor and profile.get("advisor"):
                            existing_dw.advisor = profile["advisor"]
                        if not existing_dw.overall_gpa and profile.get("overall_gpa"):
                            existing_dw.overall_gpa = profile["overall_gpa"]
                        if not existing_dw.total_credits_earned and profile.get("total_credits_earned"):
                            existing_dw.total_credits_earned = profile["total_credits_earned"]
                        if not existing_dw.degree_program and profile.get("degree_program"):
                            existing_dw.degree_program = profile["degree_program"]
                    except Exception as e:
                        print(f"[BANNER] Profile parse error: {e}")

                existing_dw.data_source = "banner_scrape"
                existing_dw.updated_at = datetime.now(timezone.utc)

                # Auto-populate user profile
                name = existing_dw.student_name
                sid = existing_dw.student_id
                if name:
                    db_user.name = name
                if sid:
                    db_user.student_id = sid

                db_user.morgan_connected = True
                db_user.morgan_connected_at = datetime.now(timezone.utc)

                sync_db.commit()

                # Count courses
                completed_count = len(json.loads(existing_dw.courses_completed or "[]"))
                ip_count = len(json.loads(existing_dw.courses_in_progress or "[]"))

                summary = {
                    "profile": bool(name),
                    "name": name or "",
                    "student_id": sid or "",
                    "classification": existing_dw.classification or "",
                    "cumulative_gpa": existing_dw.overall_gpa,
                    "total_credits": existing_dw.total_credits_earned or 0,
                    "major": existing_dw.degree_program or "",
                    "advisor": existing_dw.advisor or "",
                    "courses_completed": completed_count,
                    "courses_in_progress": ip_count,
                    "degreeworks_synced": bool(dw_json),
                    "profile_synced": bool(profile_html and len(profile_html) > 1000),
                }

                yield f"data: {json.dumps({'type': 'done', 'summary': summary})}\n\n"

            finally:
                sync_db.close()

        except ValueError as e:
            # Auth errors (safe to show: "Invalid credentials", "LDAP not available", etc.)
            yield f"data: {json.dumps({'type': 'error', 'detail': str(e)[:200]})}\n\n"
        except Exception as e:
            print(f"[ERROR] Banner sync failed: {e}")
            yield f"data: {json.dumps({'type': 'error', 'detail': 'Sync failed. Please try again.'})}\n\n"

    return StreamingResponse(generate_sse(), media_type="text/event-stream")


@app.get("/api/banner/data")
async def get_banner_data(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Returns all stored Banner data for the authenticated user."""
    banner = db.query(BannerStudentData).filter(BannerStudentData.user_id == user["user_id"]).first()

    if not banner:
        return {"connected": False, "data": None}

    return {
        "connected": True,
        "data": {
            "student_phone": banner.student_phone,
            "student_address": json.loads(banner.student_address) if banner.student_address else None,
            "current_term": banner.current_term,
            "registered_courses": json.loads(banner.registered_courses) if banner.registered_courses else [],
            "total_registered_credits": banner.total_registered_credits,
            "registration_history": json.loads(banner.registration_history) if banner.registration_history else [],
            "grade_history": json.loads(banner.grade_history) if banner.grade_history else [],
            "cumulative_gpa": banner.cumulative_gpa,
            "total_credits_earned": banner.total_credits_earned,
            "total_credits_attempted": banner.total_credits_attempted,
            "deans_list_terms": json.loads(banner.deans_list_terms) if banner.deans_list_terms else [],
            "synced_at": banner.synced_at.isoformat() if banner.synced_at else None,
            "updated_at": banner.updated_at.isoformat() if banner.updated_at else None,
        }
    }


def extract_file_content(filepath: str) -> str:
    """Reads text from PDF, DOCX, or TXT files."""
    ext = filepath.split('.')[-1].lower()
    text = ""
    try:
        if ext == 'pdf':
            #  UPDATED: Uses pypdf instead of PyPDF2
            reader = pypdf.PdfReader(filepath)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        elif ext in ['docx', 'doc']:
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            return "[Image or unsupported file type - Text extraction skipped]"
    except Exception as e:
        print(f"Error reading file: {e}")
        return f"[Error reading file content: {e}]"
    
    # Limit content to ~15k chars to fit context window
    return text[:15000]

# ==============================================================================
# Canvas LMS Integration Endpoints
# ==============================================================================

class CanvasSyncRequest(BaseModel):
    username: str
    password: str

_canvas_sync_timestamps: dict[int, list] = {}

@app.post("/api/canvas/sync")
async def sync_canvas_data(
    req: CanvasSyncRequest,
    user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Sync student data from Canvas LMS via LDAP auth. Returns SSE stream."""
    user_id = user["user_id"]

    # Rate limit: max 3 syncs per hour
    now_ts = datetime.now(timezone.utc).timestamp()
    timestamps = _canvas_sync_timestamps.get(user_id, [])
    timestamps = [t for t in timestamps if now_ts - t < 3600]
    if len(timestamps) >= 3:
        raise HTTPException(status_code=429, detail="Rate limit: max 3 Canvas syncs per hour")
    timestamps.append(now_ts)
    _canvas_sync_timestamps[user_id] = timestamps

    async def generate_sse():
        try:
            from canvas_client import sync_canvas

            progress_messages = []
            async def progress_cb(msg):
                progress_messages.append(msg)
                yield f"data: {json.dumps({'type': 'progress', 'detail': msg})}\n\n"

            # Run sync with progress streaming
            gen = progress_cb  # We need a different pattern for SSE

            yield f"data: {json.dumps({'type': 'progress', 'detail': 'Logging into Canvas...'})}\n\n"

            from canvas_client import canvas_authenticate, fetch_canvas_data

            try:
                client = await canvas_authenticate(req.username, req.password)
            except ValueError as e:
                yield f"data: {json.dumps({'type': 'error', 'detail': str(e)})}\n\n"
                return

            yield f"data: {json.dumps({'type': 'progress', 'detail': 'Fetching courses...'})}\n\n"

            try:
                data = await fetch_canvas_data(client)
            except Exception as e:
                print(f"[ERROR] Canvas fetch failed: {e}")
                yield f"data: {json.dumps({'type': 'error', 'detail': 'Failed to fetch Canvas data. Please try again.'})}\n\n"
                await client.aclose()
                return

            await client.aclose()

            yield f"data: {json.dumps({'type': 'progress', 'detail': 'Saving to database...'})}\n\n"

            # Merge grades into courses
            courses_with_grades = []
            for c in data.get("courses", []):
                grade_info = data.get("grades", {}).get(c["id"], {})
                courses_with_grades.append({
                    **c,
                    "current_score": grade_info.get("current_score"),
                    "current_grade": grade_info.get("current_grade"),
                })

            # Save to database
            try:
                existing = db.query(CanvasStudentData).filter(CanvasStudentData.user_id == user_id).first()
                if existing:
                    existing.canvas_user_id = data["profile"].get("canvas_id")
                    existing.canvas_login_id = data["profile"].get("login_id")
                    existing.courses = json.dumps(courses_with_grades)
                    existing.upcoming_assignments = json.dumps(data.get("assignments", []))
                    existing.missing_assignments = json.dumps(data.get("missing", []))
                    existing.grades = json.dumps(data.get("grades", {}))
                    existing.gradebook = json.dumps(data.get("gradebook", {}))
                    existing.updated_at = datetime.now(timezone.utc)
                else:
                    canvas_record = CanvasStudentData(
                        user_id=user_id,
                        canvas_user_id=data["profile"].get("canvas_id"),
                        canvas_login_id=data["profile"].get("login_id"),
                        courses=json.dumps(courses_with_grades),
                        upcoming_assignments=json.dumps(data.get("assignments", [])),
                        missing_assignments=json.dumps(data.get("missing", [])),
                        grades=json.dumps(data.get("grades", {})),
                        gradebook=json.dumps(data.get("gradebook", {})),
                    )
                    db.add(canvas_record)
                db.commit()
            except Exception as e:
                print(f"[ERROR] Canvas DB save failed: {e}")
                yield f"data: {json.dumps({'type': 'error', 'detail': 'Failed to save Canvas data. Please try again.'})}\n\n"
                return

            # Build summary
            summary = {
                "courses_count": len(courses_with_grades),
                "upcoming_count": len(data.get("assignments", [])),
                "missing_count": len(data.get("missing", [])),
                "courses": courses_with_grades,
                "name": data["profile"].get("name"),
                "login_id": data["profile"].get("login_id"),
            }

            yield f"data: {json.dumps({'type': 'done', 'summary': summary})}\n\n"

        except Exception as e:
            print(f"[ERROR] Canvas sync failed: {e}")
            yield f"data: {json.dumps({'type': 'error', 'detail': 'Canvas sync failed. Please try again.'})}\n\n"

    return StreamingResponse(
        generate_sse(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"}
    )



@app.get("/api/canvas")
async def get_canvas_data(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Get stored Canvas data for the current user."""
    canvas = db.query(CanvasStudentData).filter(CanvasStudentData.user_id == user["user_id"]).first()
    if not canvas:
        return {"connected": False}

    return {
        "connected": True,
        "canvas_login_id": canvas.canvas_login_id,
        "courses": json.loads(canvas.courses) if canvas.courses else [],
        "upcoming_assignments": json.loads(canvas.upcoming_assignments) if canvas.upcoming_assignments else [],
        "missing_assignments": json.loads(canvas.missing_assignments) if canvas.missing_assignments else [],
        "grades": json.loads(canvas.grades) if canvas.grades else {},
        "gradebook": json.loads(canvas.gradebook) if canvas.gradebook else {},
        "synced_at": canvas.synced_at.isoformat() if canvas.synced_at else None,
        "updated_at": canvas.updated_at.isoformat() if canvas.updated_at else None,
    }


@app.delete("/api/canvas/disconnect")
async def disconnect_canvas(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Remove Canvas data for the current user."""
    canvas = db.query(CanvasStudentData).filter(CanvasStudentData.user_id == user["user_id"]).first()
    if canvas:
        db.delete(canvas)
        db.commit()
    return {"success": True, "message": "Canvas disconnected"}


# ==============================================================================
# MOMENTUM SCORE - Academic Performance Index
# ==============================================================================
from services.canvas_analytics import compute_momentum_score

@app.get("/api/momentum-score")
async def momentum_score(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Compute academic momentum score from Canvas + DegreeWorks + Banner data."""
    canvas = db.query(CanvasStudentData).filter(CanvasStudentData.user_id == user["user_id"]).first()
    dw = db.query(DegreeWorksData).filter(DegreeWorksData.user_id == user["user_id"]).first()
    banner = db.query(BannerStudentData).filter(BannerStudentData.user_id == user["user_id"]).first()

    canvas_dict = {
        "courses": canvas.courses,
        "gradebook": canvas.gradebook,
        "missing_assignments": canvas.missing_assignments,
    } if canvas else None

    dw_dict = {
        "overall_gpa": dw.overall_gpa,
        "total_credits_earned": dw.total_credits_earned,
        "credits_required": dw.credits_required,
        "classification": dw.classification,
    } if dw else None

    banner_dict = {
        "cumulative_gpa": banner.cumulative_gpa,
    } if banner else None

    return compute_momentum_score(canvas_dict, dw_dict, banner_dict)


# ==============================================================================
# RIPPLE EFFECT - Prerequisite Dependency Graph
# ==============================================================================
from services.prereq_engine import build_prerequisite_graph

@app.get("/api/ripple-effect")
async def ripple_effect(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Get prerequisite dependency graph with student status overlay."""
    dw_dict = await asyncio.to_thread(_fetch_dw_sync, user["user_id"])
    canvas_dict = await asyncio.to_thread(_fetch_canvas_sync, user["user_id"])
    return build_prerequisite_graph(dw_dict, canvas_dict)


# ==============================================================================
# GRADE SURGEON - Canvas Grade Analysis
# ==============================================================================
from services.canvas_analytics import analyze_course_grade, get_all_courses_summary, parse_gradebook

@app.get("/api/grade-analysis")
async def grade_analysis_all(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Get grade analysis summary for all courses."""
    canvas = db.query(CanvasStudentData).filter(CanvasStudentData.user_id == user["user_id"]).first()
    if not canvas or not canvas.gradebook:
        raise HTTPException(404, "No gradebook data. Please sync Canvas first.")
    gradebook = parse_gradebook(canvas.gradebook)
    courses = json.loads(canvas.courses) if canvas.courses else []
    return {
        "courses": get_all_courses_summary(gradebook, courses),
        "synced_at": canvas.synced_at.isoformat() if canvas.synced_at else None,
        "updated_at": canvas.updated_at.isoformat() if canvas.updated_at else None,
    }

@app.get("/api/grade-analysis/{course_id}")
async def grade_analysis_course(course_id: str, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Get detailed grade analysis for a specific course."""
    canvas = db.query(CanvasStudentData).filter(CanvasStudentData.user_id == user["user_id"]).first()
    if not canvas or not canvas.gradebook:
        raise HTTPException(404, "No gradebook data. Please sync Canvas first.")
    gradebook = parse_gradebook(canvas.gradebook)
    if course_id not in gradebook:
        raise HTTPException(404, f"Course {course_id} not found in gradebook.")
    courses = json.loads(canvas.courses) if canvas.courses else []
    course_name = next((c.get("name", "") for c in courses if str(c.get("id", "")) == course_id), "Unknown")
    return analyze_course_grade(gradebook[course_id], course_name)


# ==============================================================================
# PARALLEL DB HELPERS (Thread-safe, each creates its own session)
# ==============================================================================

def _fetch_dw_sync(user_id: int) -> Optional[dict]:
    """Fetch DegreeWorks + Banner data in a separate DB session for parallel execution."""
    db = SessionLocal()
    try:
        dw = db.query(DegreeWorksData).filter(DegreeWorksData.user_id == user_id).first()
        if not dw:
            return None
        result = {
            "student_name": dw.student_name,
            "student_id": dw.student_id,
            "classification": dw.classification,
            "degree_program": dw.degree_program,
            "overall_gpa": dw.overall_gpa,
            "major_gpa": dw.major_gpa,
            "total_credits_earned": dw.total_credits_earned,
            "credits_required": dw.credits_required,
            "credits_remaining": dw.credits_remaining,
            "advisor": dw.advisor,
            "catalog_year": dw.catalog_year,
            "courses_completed": dw.courses_completed,
            "courses_in_progress": dw.courses_in_progress,
            "courses_remaining": dw.courses_remaining,
            "raw_data": dw.raw_data,
            "data_source": getattr(dw, 'data_source', None) or "manual_entry",
        }

        # Also fetch Banner data if available
        banner = db.query(BannerStudentData).filter(BannerStudentData.user_id == user_id).first()
        if banner:
            result["banner"] = {
                "current_term": banner.current_term,
                "registered_courses": banner.registered_courses,
                "total_registered_credits": banner.total_registered_credits,
                "registration_history": banner.registration_history,
                "grade_history": banner.grade_history,
                "cumulative_gpa": banner.cumulative_gpa,
                "total_credits_earned": banner.total_credits_earned,
                "total_credits_attempted": banner.total_credits_attempted,
                "deans_list_terms": banner.deans_list_terms,
            }

        return result
    finally:
        db.close()


def _fetch_canvas_sync(user_id: int) -> Optional[dict]:
    """Fetch Canvas LMS data in a separate DB session for parallel execution."""
    db = SessionLocal()
    try:
        canvas = db.query(CanvasStudentData).filter(CanvasStudentData.user_id == user_id).first()
        if not canvas:
            return None
        return {
            "courses": canvas.courses,
            "upcoming_assignments": canvas.upcoming_assignments,
            "missing_assignments": canvas.missing_assignments,
            "grades": canvas.grades,
            "gradebook": canvas.gradebook,
            "synced_at": str(canvas.synced_at) if canvas.synced_at else None,
            "updated_at": str(canvas.updated_at) if canvas.updated_at else None,
        }
    finally:
        db.close()


# Context builders extracted to services/context_builders.py
from services.context_builders import (
    sanitize_canvas_field as _sanitize_canvas_field,
    format_short_date as _format_short_date,
    build_canvas_context as _build_canvas_context,
)

# Tier 1: Query rewriting for follow-up resolution
from services.query_rewriter import rewrite_query, is_likely_followup

# Tier 2: Long-term user memory
from services.memory_service import fetch_user_memories_sync, build_memory_context
from services.course_context import build_course_context


def _fetch_history_sync(user_id: int, session_id: str, limit: int = 10) -> list:
    """Fetch chat history in a separate DB session for parallel execution."""
    db = SessionLocal()
    try:
        history = db.query(ChatHistory)\
            .filter(ChatHistory.user_id == user_id, ChatHistory.session_id == session_id)\
            .order_by(ChatHistory.timestamp.desc())\
            .limit(limit)\
            .all()
        return [{"user_query": h.user_query, "bot_response": h.bot_response} for h in reversed(history)]
    finally:
        db.close()


from services.context_builders import (
    build_student_context as _build_student_context,
    build_conversation_context as _build_conversation_context,
)


# --- CHAT ROUTES (WITH CONVERSATION MEMORY + PERSONALIZATION) ---
@app.post("/chat")
async def chat_with_bot(req: QueryRequest, user=Depends(get_current_user), db: Session = Depends(get_db)):
    if not user: raise HTTPException(401, "Unauthorized")

    user_q = req.query.strip()
    original_q = user_q  # Preserve original for chat history (before rewrite)
    session_id = req.session_id or "default"

    # Detect file upload early to decide what data we need
    file_match = re.search(r'uploads/chat_files/([^\)]+)', user_q)
    # Always fetch history for follow-up rewriting (Tier 1) + file uploads + legacy path
    needs_history = True

    # Lazy-load: only fetch Canvas if query mentions grades/assignments/deadlines/course codes
    CANVAS_KEYWORDS = {"grade", "assignment", "due", "deadline", "missing", "class",
                       "course", "score", "submit", "canvas", "homework", "quiz",
                       "test", "exam", "gpa", "taking", "enrolled"}
    has_course_code = bool(re.search(r'\b[A-Z]{2,4}\s*\d{3}\b', user_q, re.IGNORECASE))
    needs_canvas = has_course_code or any(kw in user_q.lower() for kw in CANVAS_KEYWORDS)

    # Parallel fetch: DegreeWorks + Canvas (if needed) + chat history (for rewriting) + long-term memory
    fetch_tasks = [
        asyncio.to_thread(_fetch_dw_sync, user["user_id"]),
        asyncio.to_thread(_fetch_history_sync, user["user_id"], session_id, 5),
        asyncio.to_thread(fetch_user_memories_sync, user["user_id"], 10),
    ]
    if needs_canvas:
        fetch_tasks.append(asyncio.to_thread(_fetch_canvas_sync, user["user_id"]))

    results = await asyncio.gather(*fetch_tasks, return_exceptions=True)

    dw_dict = results[0] if not isinstance(results[0], Exception) else None
    history_dicts = results[1] if not isinstance(results[1], Exception) else []
    memory_dicts = results[2] if not isinstance(results[2], Exception) else []
    canvas_dict = results[3] if needs_canvas and len(results) > 3 and not isinstance(results[3], Exception) else None

    # Tier 1: Rewrite follow-up queries to be self-contained (fixes pronoun resolution)
    if USE_VERTEX_AGENT and history_dicts and is_likely_followup(user_q):
        user_q = await asyncio.to_thread(rewrite_query, user_q, history_dicts)

    student_context = _build_student_context(dw_dict) if dw_dict else ""
    canvas_context = _build_canvas_context(canvas_dict) if canvas_dict else ""
    memory_context = build_memory_context(memory_dicts)
    conversation_context = _build_conversation_context(history_dicts)

    # Pre-compute course context (prereq analysis, schedule, eligibility)
    course_context = build_course_context(dw_dict, user_q) if dw_dict else ""
    if course_context:
        # Append to student context so agent sees it alongside DegreeWorks data
        student_context += f"\n{course_context}"

    if file_match and USE_VERTEX_AGENT:
        # File uploaded -> include file content as context for the agent
        filename = file_match.group(1)
        filepath = os.path.join(CHAT_FILES_FOLDER, filename)

        if os.path.exists(filepath):
            file_content = extract_file_content(filepath)
            clean_query = re.sub(r'\[.*?\]\(.*?\)', '', user_q).strip()
            if not clean_query: clean_query = "Summarize this file."

            file_context = f"{student_context}{canvas_context}{conversation_context}File Content:\n{file_content}\n"
            answer = query_agent(
                query=clean_query,
                user_id=str(user["user_id"]),
                context=file_context,
                model=req.model,
                canvas_context=canvas_context,
                memory_context=memory_context,
            )
        else:
            answer = "I received the file link, but I cannot find the file on the server to read it."

    elif file_match and llm:
        # Legacy: File uploaded with old LLM pipeline
        filename = file_match.group(1)
        filepath = os.path.join(CHAT_FILES_FOLDER, filename)

        if os.path.exists(filepath):
            file_content = extract_file_content(filepath)
            system_msg = f"""You are a helpful academic assistant for Morgan State University's Computer Science department.
Use the provided file content and conversation history to answer the user's question.
{student_context}"""

            clean_query = re.sub(r'\[.*?\]\(.*?\)', '', user_q).strip()
            if not clean_query: clean_query = "Summarize this file."

            user_msg = f"{conversation_context}File Content:\n{file_content}\n\nCurrent Question: {clean_query}"

            try:
                response = llm([
                    SystemMessage(content=system_msg),
                    HumanMessage(content=user_msg)
                ])
                answer = response.content
            except Exception as e:
                answer = f"I read the file, but had trouble analyzing it: {e}"
        else:
            answer = "I received the file link, but I cannot find the file on the server to read it."

    elif USE_VERTEX_AGENT:
        # Vertex AI Agent Engine path
        # Tier 1: Query already rewritten above (follow-ups resolved)
        # Tier 2: Long-term memory injected via memory_context
        # NOTE: DegreeWorks = stable context (hashed for session reuse)
        #       Canvas + Memory = volatile (sent via state_delta per request)
        try:
            agent_context = student_context  # DegreeWorks only (stable, for session reuse)

            print(f" Vertex AI query: '{user_q[:50]}...' (user={user['user_id']}, context={len(agent_context)} chars, memory={len(memory_context)} chars, model={req.model})")
            answer = query_agent(
                query=user_q,
                user_id=str(user["user_id"]),
                context=agent_context,
                model=req.model,
                canvas_context=canvas_context,
                memory_context=memory_context,
            )
        except Exception as e:
            print(f"   Vertex AI Chat Error: {e}")
            answer = "I'm having trouble processing your request. Please try again."
    elif llm and retriever:
        # Legacy Pinecone + OpenAI RAG path (fallback)
        norm = re.sub(r'[\s\W]+', '', user_q.lower())
        if re.match(r'^(hi|hello|hey)\b', user_q.lower()):
            answer = "Hello! How can I help you today?"
        elif re.match(r'^(bye|goodbye|see you)\b', user_q.lower()):
            answer = "Goodbye! Have a great day."
        elif re.search(r'\b(thankyou|thanks|thanx|thx|ty)\b', norm):
            answer = "You're welcome! Let me know if you have any other questions."
        else:
            try:
                docs = retriever.get_relevant_documents(user_q)
                context_docs = "\n\n".join([doc.page_content for doc in docs[:8]])
                system_prompt = f"""You are CS Navigator, an academic assistant for Morgan State University's CS department.
{student_context}
ONLY answer based on the KNOWLEDGE BASE CONTEXT provided. If info is not found, say so honestly."""
                full_message = f"{conversation_context}Knowledge base:\n{context_docs}\n\nQuestion: {user_q}"
                response = llm([
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=full_message)
                ])
                answer = response.content.strip()
            except Exception as e:
                print(f"   Legacy Chat Error: {e}")
                answer = "I'm having trouble processing your request."
    else:
        answer = "AI system is initializing. Please try again in a moment."

    # 3. SAVE to RDS (User-Specific)
    try:
        new_chat = ChatHistory(
            user_id=user["user_id"],
            session_id=session_id,
            user_query=original_q,
            bot_response=answer
        )
        db.add(new_chat)
        db.commit()
    except Exception as e:
        print(f"[ERROR] Failed to save chat history: {e}")

    # 4. Track failed queries for auto-research agent
    if answer and "error" not in answer.lower()[:50]:
        try:
            from research_agent import detect_and_log_failed_query
            detect_and_log_failed_query(original_q, answer, user["user_id"], has_student_data=bool(student_context))
        except Exception:
            pass

    return {"response": answer}


# ==============================================================================
# STREAMING CHAT ENDPOINT (Server-Sent Events)
# ==============================================================================
@app.post("/chat/stream")
async def chat_stream(req: QueryRequest, user=Depends(get_current_user), db: Session = Depends(get_db)):
    """
    Streaming chat endpoint using Server-Sent Events (SSE).
    Returns text chunks as they arrive from the AI agent for faster perceived response time.

    v4.2: Uses async parallel DB fetch and shared _build_student_context helper.
    Chat history fetch removed (not used in Vertex path, ADK manages its own memory).
    """
    if not user:
        raise HTTPException(401, "Unauthorized")

    user_q = req.query.strip()
    original_q = user_q  # Keep original for cache key + chat history
    session_id = req.session_id or "default"
    user_id = user["user_id"]

    # Lazy-load: only fetch Canvas if query mentions grades/assignments/deadlines/course codes
    CANVAS_KEYWORDS = {"grade", "assignment", "due", "deadline", "missing", "class",
                       "course", "score", "submit", "canvas", "homework", "quiz",
                       "test", "exam", "gpa", "taking", "enrolled"}
    has_course_code = bool(re.search(r'\b[A-Z]{2,4}\s*\d{3}\b', user_q, re.IGNORECASE))
    needs_canvas = has_course_code or any(kw in user_q.lower() for kw in CANVAS_KEYWORDS)

    # Non-blocking parallel fetch: DegreeWorks + Canvas (if needed) + history (for rewriting) + memory
    fetch_tasks = [
        asyncio.to_thread(_fetch_dw_sync, user_id),
        asyncio.to_thread(_fetch_history_sync, user_id, session_id, 5),
        asyncio.to_thread(fetch_user_memories_sync, user_id, 10),
    ]
    if needs_canvas:
        fetch_tasks.append(asyncio.to_thread(_fetch_canvas_sync, user_id))

    results = await asyncio.gather(*fetch_tasks, return_exceptions=True)

    dw_dict = results[0] if not isinstance(results[0], Exception) else None
    history_dicts = results[1] if not isinstance(results[1], Exception) else []
    memory_dicts = results[2] if not isinstance(results[2], Exception) else []
    canvas_dict = results[3] if needs_canvas and len(results) > 3 and not isinstance(results[3], Exception) else None

    # Tier 1: Rewrite follow-up queries (resolve pronouns before KB search)
    if history_dicts and is_likely_followup(user_q):
        user_q = await asyncio.to_thread(rewrite_query, user_q, history_dicts)

    student_context = _build_student_context(dw_dict) if dw_dict else ""
    canvas_context = _build_canvas_context(canvas_dict) if canvas_dict else ""
    memory_context = build_memory_context(memory_dicts)

    # Pre-compute course context (prereq analysis, schedule, eligibility)
    course_context = build_course_context(dw_dict, user_q) if dw_dict else ""
    if course_context:
        student_context += f"\n{course_context}"

    agent_context = student_context  # DegreeWorks + course analysis (stable, for session reuse)

    # =========================================================================
    # CACHE CHECK - Return cached response instantly if available
    # =========================================================================
    context_hash = get_context_hash(user_id, has_degreeworks=bool(dw_dict), model=req.model, has_canvas=bool(canvas_dict))

    # Skip cache when user taps "Regenerate" for a fresh answer
    if req.skip_cache:
        print(f"[CACHE] SKIP (regenerate) for query: {user_q[:50]}...")
        cached_response = None
        # Force new ADK session so agent re-queries the search index fresh
        import time as _time
        context_hash = f"regen_{int(_time.time())}"
        reset_session(str(user_id))
    else:
        cached_response = query_cache.get(user_q, context_hash)

    if cached_response:
        print(f"[CACHE] HIT for query: {user_q[:50]}...")

        async def generate_cached_sse():
            """Return cached response as SSE."""
            # Send status to show it's from cache
            yield f"data: {json.dumps({'type': 'status', 'content': 'Retrieved from cache'})}\n\n"
            # Send the full response immediately
            yield f"data: {json.dumps({'type': 'done', 'content': cached_response})}\n\n"

            # Still save to chat history (save original query, not rewritten)
            try:
                with SessionLocal() as save_db:
                    new_chat = ChatHistory(
                        user_id=user_id,
                        session_id=session_id,
                        user_query=original_q,
                        bot_response=cached_response
                    )
                    save_db.add(new_chat)
                    save_db.commit()
            except Exception as e:
                print(f"[ERROR] Failed to save cached chat history: {e}")

        return StreamingResponse(
            generate_cached_sse(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no"
            }
        )

    # =========================================================================
    # CACHE MISS - Stream from AI agent and cache the result
    # =========================================================================
    print(f"[CACHE] MISS for query: {user_q[:50]}...")
    stream_had_error = False

    async def generate_sse():
        """SSE generator that streams text chunks from the agent."""
        nonlocal stream_had_error
        full_response = ""
        try:
            for event in query_agent_stream(
                query=user_q,
                user_id=str(user_id),
                context=agent_context,
                model=req.model,
                canvas_context=canvas_context,
                memory_context=memory_context,
            ):
                event_type = event.get("type", "")
                content = event.get("content", "")

                if event_type == "status":
                    yield f"data: {json.dumps({'type': 'status', 'content': content})}\n\n"

                elif event_type == "chunk":
                    full_response += content
                    yield f"data: {json.dumps({'type': 'chunk', 'content': content})}\n\n"

                elif event_type == "done":
                    full_response = content or full_response
                    yield f"data: {json.dumps({'type': 'done', 'content': full_response})}\n\n"

                elif event_type == "error":
                    stream_had_error = True
                    yield f"data: {json.dumps({'type': 'error', 'content': content})}\n\n"
                    full_response = content
                    break

        except Exception as e:
            stream_had_error = True
            print(f"[ERROR] Streaming error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'content': 'An error occurred during streaming.'})}\n\n"
            full_response = "An error occurred during streaming."

        # Cache the successful response
        if full_response and "error" not in full_response.lower()[:50]:
            if query_cache.set(user_q, full_response, context_hash):
                print(f"[CACHE] Stored response for: {user_q[:50]}...")

        # Save to chat history after stream completes (save original query, not rewritten)
        try:
            with SessionLocal() as save_db:
                new_chat = ChatHistory(
                    user_id=user_id,
                    session_id=session_id,
                    user_query=original_q,
                    bot_response=full_response
                )
                save_db.add(new_chat)
                save_db.commit()
        except Exception as e:
            print(f"[ERROR] Failed to save streamed chat history: {e}")

        # Track failed queries for auto-research agent
        # Skip detection on error/empty responses (infra errors aren't KB misses)
        if full_response and not stream_had_error and "error" not in full_response.lower()[:50]:
            try:
                from research_agent import detect_and_log_failed_query
                detect_and_log_failed_query(original_q, full_response, user_id, has_student_data=bool(agent_context))
            except Exception:
                pass

    return StreamingResponse(
        generate_sse(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


# ==============================================================================
# GUEST CHAT ENDPOINT (No Authentication Required)
# ==============================================================================
@app.post("/chat/guest")
async def chat_guest(req: GuestQueryRequest, request: Request):
    """
    Guest chat endpoint - NO authentication required.
    - No personalization (no DegreeWorks)
    - No history persistence
    - Rate limited: 10 requests/minute per IP
    """
    # Get client IP for rate limiting
    client_ip = request.client.host if request.client else "unknown"

    # Check rate limit
    if not check_guest_rate_limit(client_ip):
        raise HTTPException(
            status_code=429,
            detail="Rate limit exceeded. Please try again in a minute or sign up for unlimited access!"
        )

    user_q = req.query.strip()
    if not user_q:
        return {"response": "Please enter a question."}

    # #11 - Limit query length (500 chars max)
    if len(user_q) > 500:
        user_q = user_q[:500]

    # Small talk override - handle greetings, acknowledgments, and non-questions
    lower_q = user_q.lower().strip()
    norm = re.sub(r'[\s\W]+', '', lower_q)
    word_count = len(lower_q.split())

    # #9 FIX: Only match greetings if it's JUST a greeting (1-2 words max)
    # Greetings (including typos) - only if short message
    greeting_patterns = ['hi', 'hey', 'heyt', 'hii', 'heyy', 'hello', 'helo', 'howdy', 'sup', 'yo', 'hola', 'greetings']
    if word_count <= 2 and (norm in greeting_patterns or re.match(r'^(hi+|hey+t?|hello+)$', norm)):
        return {"response": "Hello! I'm CS Navigator. How can I help you learn about Morgan State's Computer Science program today?"}

    # #8 FIX: "what's up", "how are you" patterns
    elif norm in ['whatsup', 'wassup', 'wazzup', 'whatsgood', 'howareyou', 'howru', 'howreyou', 'howyoudoing']:
        return {"response": "I'm doing great, thanks for asking! How can I help you with Morgan State's CS program today?"}

    # Goodbyes - only if short
    elif word_count <= 3 and re.match(r'^(bye|goodbye|see you|later|cya|peace|gotta go|gtg)', lower_q):
        return {"response": "Goodbye! Sign up for a free account to save your chat history and get personalized advice!"}

    # Thank you
    elif re.search(r'\b(thank|thanks|thanx|thx|ty|appreciate)\b', lower_q):
        return {"response": "You're welcome! Feel free to ask more questions. Sign up to unlock personalized features!"}

    # #8 FIX: Reactions and fillers (lol, haha, test, etc.)
    elif norm in ['lol', 'lmao', 'rofl', 'haha', 'hahaha', 'hehe', 'lolol', 'xd', 'test', 'testing', 'testtest', 'asdf', 'aaa', 'zzz', 'idk', 'idc', 'nvm', 'nevermind', 'bruh', 'bro', 'dude', 'wow', 'omg', 'wtf', 'wth']:
        return {"response": "I'm here whenever you're ready! Ask me anything about Morgan State's CS program - courses, professors, requirements, or career paths."}

    # Acknowledgments (ok, sure, cool, etc.)
    elif norm in ['ok', 'okay', 'okk', 'okok', 'k', 'kk', 'sure', 'alright', 'aight', 'cool', 'nice', 'great', 'good', 'gotit', 'understood', 'isee', 'ah', 'oh', 'ohh', 'hmm', 'hm', 'mhm', 'yep', 'yup', 'yes', 'yeah', 'ya', 'no', 'nope', 'nah', 'fine', 'bet', 'word', 'facts', 'true', 'right', 'correct']:
        return {"response": "Got it! Feel free to ask me anything about Morgan State's CS program - courses, professors, requirements, or career opportunities!"}

    # Very short inputs (1-2 chars) or just punctuation/emojis
    elif len(norm) <= 2 or not any(c.isalpha() for c in user_q):
        return {"response": "I'm here to help! Ask me about CS courses, professors, degree requirements, or anything else about Morgan State's Computer Science program."}

    # =========================================================================
    # CACHE CHECK - Return cached response instantly for guest queries
    # =========================================================================
    # Guest queries share cache (no user-specific context)
    cached_response = query_cache.get(user_q, context_hash="")
    if cached_response:
        print(f"[CACHE] HIT (guest) for: {user_q[:50]}...")
        return {"response": cached_response, "cached": True}

    # Detect personal academic queries and redirect guests to sign up
    _PERSONAL_KEYWORDS = [
        "my gpa", "my grade", "my classes", "my schedule", "my advisor",
        "my courses", "my transcript", "my degree", "my credits",
        "my assignment", "my canvas", "degreeworks", "degree works",
        "what am i taking", "what classes am i", "how many credits do i",
        "my remaining", "my progress", "my academic"
    ]
    query_lower = user_q.lower()
    if any(kw in query_lower for kw in _PERSONAL_KEYWORDS):
        return {"response": (
            "To access your personal academic information like GPA, courses, "
            "and degree progress, you'll need to **create a free account** with your "
            "Morgan State email. This connects your DegreeWorks and Canvas data securely.\n\n"
            "**[Create an account here](https://cs.inavigator.ai/register)** to unlock personalized features!"
        )}

    # Use Vertex AI Agent for real questions
    if USE_VERTEX_AGENT:
        try:
            # Use a unique guest_user_id per request to prevent session bleed.
            # Previously IP-based, which caused students on the same campus WiFi
            # to share ADK sessions and see each other's DegreeWorks data.
            import uuid
            guest_user_id = f"guest_{uuid.uuid4().hex[:12]}"
            print(f"[CACHE] MISS (guest) for: '{user_q[:50]}...'")
            answer = query_agent(
                query=user_q,
                user_id=guest_user_id,
                context="",
            )

            # Cache the successful response
            if answer and "error" not in answer.lower()[:50]:
                query_cache.set(user_q, answer, context_hash="")

        except Exception as e:
            print(f"   Guest Vertex AI Error: {e}")
            answer = "I'm having trouble processing your request. Please try again."
    elif llm and retriever:
        # Legacy Pinecone + OpenAI RAG path (fallback)
        try:
            docs = retriever.get_relevant_documents(user_q)
            context_docs = "\n\n".join([doc.page_content for doc in docs[:8]])
            if not context_docs.strip():
                answer = "I don't have specific information about that. Contact the CS department at compsci@morgan.edu or (443) 885-3962."
            else:
                response = llm([
                    SystemMessage(content="You are CS Navigator for Morgan State University's CS department. ONLY answer from the provided context."),
                    HumanMessage(content=f"Context:\n{context_docs}\n\nQuestion: {user_q}")
                ])
                answer = response.content.strip()
        except Exception as e:
            print(f"   Guest Legacy Error: {e}")
            answer = "I'm having trouble processing your request. Please try again."
    else:
        answer = "AI system is initializing. Please try again in a moment."

    # Track failed queries for auto-research agent (guest queries too)
    if answer and "error" not in answer.lower()[:50]:
        try:
            from research_agent import detect_and_log_failed_query
            detect_and_log_failed_query(user_q, answer)
        except Exception:
            pass

    return {"response": answer}

@app.get("/chat-history")
async def get_chat_history(user=Depends(get_current_user), db: Session = Depends(get_db)):
    """Fetch chat history for the logged-in user from RDS"""
    chats = db.query(ChatHistory)\
              .filter(ChatHistory.user_id == user["user_id"])\
              .order_by(ChatHistory.timestamp.asc())\
              .all()
    
    # Format for frontend
    history = []
    for c in chats:
        history.append({
            "session_id": c.session_id or "default",
            "user": c.user_query,
            "bot": c.bot_response,
            "time": c.timestamp.isoformat()
        })
        
    return {"history": history}

@app.post("/reset-history")
async def reset_chat_history(user=Depends(get_current_user), db: Session = Depends(get_db)):
    """Delete history only for this user"""
    db.query(ChatHistory).filter(ChatHistory.user_id == user["user_id"]).delete()
    db.commit()
    return {"message": "Chat history reset."}

# --- Voice Mode Endpoints ---
@app.post("/api/tts")
async def text_to_speech(req: TTSRequest, _user=Depends(get_current_user)):
    """Convert text to speech using OpenAI TTS API"""
    if not OPENAI_API_KEY:
        raise HTTPException(500, "OpenAI API key not configured")

    try:
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)

        # Use TTS-1 for speed (tts-1-hd for quality but slower)
        response = client.audio.speech.create(
            model="tts-1",
            voice=req.voice,
            input=req.text[:4096],  # Limit to 4096 chars
            response_format="mp3"
        )

        # Stream the audio response
        audio_data = io.BytesIO(response.content)
        return StreamingResponse(
            audio_data,
            media_type="audio/mpeg",
            headers={"Content-Disposition": "inline; filename=response.mp3"}
        )
    except Exception as e:
        print(f"TTS Error: {e}")
        raise HTTPException(500, f"TTS generation failed: {str(e)}")

@app.get("/api/popular-questions")
async def get_popular_questions():
    """Returns 8 randomly selected questions from a curated pool."""
    import random

    QUESTION_POOL = [
        # Course & curriculum
        "What courses should I take next semester if I'm interested in AI/ML?",
        "Can you recommend a study plan for the cybersecurity track?",
        "What are the prerequisites for COSC 450 Operating Systems?",
        "What electives count toward the CS degree?",
        "What math courses are required for the CS major?",
        "What is the recommended course sequence for freshmen CS students?",
        "Which courses cover data structures and algorithms?",
        # Department & faculty
        "Who are the professors in the CS department and what do they teach?",
        "Who is the chair of the Computer Science department?",
        "What research areas do CS faculty specialize in?",
        "How do I find a faculty mentor for my capstone project?",
        # Career & opportunities
        "What internship and co-op opportunities are available for CS majors?",
        "What career paths can I pursue with a CS degree from Morgan State?",
        "How can I prepare for technical interviews?",
        "What companies recruit CS students from Morgan State?",
        # Academic advising & graduation
        "How do I apply for graduation and what requirements do I need?",
        "How many credits do I need to graduate with a CS degree?",
        "What is the difference between a B.S. and B.A. in Computer Science?",
        "What is the minimum GPA required to stay in the CS program?",
        # Research & extracurricular
        "What research labs and projects can I join in the CS department?",
        "Are there any CS student organizations or clubs at Morgan State?",
        "How can I get involved in undergraduate research?",
        "What programming competitions can Morgan State students participate in?",
        # Frequently asked
        "How do I contact my academic advisor?",
        "Where is the Computer Science department located?",
        "How do I register for CS courses?",
    ]

    return {"questions": random.sample(QUESTION_POOL, 8)}

# --- Admin / Ingest Routes ---
@app.post("/ingest")
async def ingest_data_endpoint(user=Depends(get_current_user)):
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admins only")

    files = [os.path.join(DATA_DIR, fn) for fn in sorted(os.listdir(DATA_DIR)) if fn.lower().endswith(".json")]
    raw = []
    for p in files:
        raw.extend(load_json_documents([p]))

    splitter = TokenTextSplitter(chunk_size=800, chunk_overlap=160, model_name="gpt-3.5-turbo")
    texts, metas = [], []
    for doc in raw:
        for chunk in splitter.split_text(doc["text"]):
            texts.append(chunk)
            metas.append({"source": os.path.basename(doc["source"])})

    embeddings = OpenAIEmbeddings(model="text-embedding-3-small", openai_api_key=OPENAI_API_KEY)
    PineconeVectorStore.from_texts(
        texts=texts,
        embedding=embeddings,
        metadatas=metas,
        index_name=PINECONE_INDEX,
        namespace=PINECONE_NAMESPACE,
    )
    return {"message": f"Ingested into {PINECONE_INDEX}:{PINECONE_NAMESPACE}", "chunks": len(texts)}

@app.delete("/clear-index")
async def clear_index(user=Depends(get_current_user)):
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admins only")
    if not pc:
        raise HTTPException(status_code=500, detail="Pinecone not initialized")
    idx = pc.Index(PINECONE_INDEX)
    idx.delete(delete_all=True, namespace=PINECONE_NAMESPACE)
    return {"message": f"Cleared namespace '{PINECONE_NAMESPACE}' in index {PINECONE_INDEX}"}

# --- Curriculum Routes ---
@app.post("/api/curriculum/add")
async def add_course(course: Course, user=Depends(get_current_user)):
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admins only")
    arr = json.load(open(CLASSES_FILE, encoding="utf-8"))
    arr.append(course.model_dump())
    json.dump(arr, open(CLASSES_FILE, "w", encoding="utf-8"), indent=2)
    return {"message": "Course added", "course": course}

@app.delete("/api/curriculum/delete/{code}")
async def delete_course(code: str, user=Depends(get_current_user)):
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admins only")
    arr = json.load(open(CLASSES_FILE, encoding="utf-8"))
    filtered = [c for c in arr if c.get("course_code") != code]
    json.dump(filtered, open(CLASSES_FILE, "w", encoding="utf-8"), indent=2)
    return {"message": f"{code} deleted"}

@app.get("/api/curriculum")
async def get_curriculum():
    """Returns full curriculum data including degree info, courses, and elective requirements.
    Source of truth: courses.txt (KB file). Falls back to classes.json if txt not available."""
    try:
        # Primary: parse from txt knowledge base (single source of truth)
        if os.path.exists(KB_COURSES_FILE):
            return parse_curriculum_from_txt()

        # Fallback: classes.json (legacy)
        data = json.load(open(CLASSES_FILE, encoding="utf-8"))

        if isinstance(data, dict) and "courses" in data:
            return {
                "degree_info": data.get("degree_info", {}),
                "courses": data.get("courses", []),
                "elective_requirements": data.get("elective_requirements", {})
            }

        if isinstance(data, list):
            return {"degree_info": {}, "courses": data, "elective_requirements": {}}

        for key in ("computer_science_courses", "classes"):
            arr = data.get(key)
            if isinstance(arr, list):
                return {"degree_info": {}, "courses": arr, "elective_requirements": {}}

        return {"degree_info": {}, "courses": [], "elective_requirements": {}}
    except FileNotFoundError:
        return {"degree_info": {}, "courses": [], "elective_requirements": {}}

@app.get("/health")
def health():
    if USE_VERTEX_AGENT:
        try:
            result = check_agent_health()
            ai_status = result.get("status", "offline") if isinstance(result, dict) else "offline"
        except Exception:
            ai_status = "offline"
        return {"status": "ok", "db": "connected", "ai": "ready" if ai_status == "connected" else "offline"}
    return {"status": "ok", "db": "connected", "ai": "ready" if qa else "offline"}

# ==============================================================================
# ADMIN DASHBOARD ENDPOINTS
# ==============================================================================

# --- Admin: User Management ---
@app.get("/api/admin/users")
async def get_all_users(
    search: Optional[str] = None,
    role: Optional[str] = None,
    user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all users (admin only)"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    query = db.query(User).order_by(User.created_at.desc())

    if search:
        search_term = f"%{search}%"
        query = query.filter(
            (User.email.ilike(search_term)) |
            (User.name.ilike(search_term)) |
            (User.student_id.ilike(search_term))
        )

    if role and role != "all":
        query = query.filter(User.role == role)

    users = query.all()

    return {
        "users": [
            {
                "id": u.id,
                "email": u.email,
                "name": u.name,
                "role": u.role,
                "student_id": u.student_id,
                "major": u.major,
                "morgan_connected": u.morgan_connected,
                "created_at": u.created_at.isoformat() if u.created_at else None
            }
            for u in users
        ],
        "total": len(users)
    }

@app.get("/api/admin/users/stats")
async def get_user_stats(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Get user statistics (admin only)"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    from datetime import timedelta
    now = datetime.now(timezone.utc)
    week_ago = now - timedelta(days=7)
    month_ago = now - timedelta(days=30)

    total_users = db.query(User).count()
    total_students = db.query(User).filter(User.role == "student").count()
    total_admins = db.query(User).filter(User.role == "admin").count()
    new_this_week = db.query(User).filter(User.created_at >= week_ago).count()
    new_this_month = db.query(User).filter(User.created_at >= month_ago).count()
    morgan_connected = db.query(User).filter(User.morgan_connected == True).count()

    return {
        "total": total_users,
        "students": total_students,
        "admins": total_admins,
        "new_this_week": new_this_week,
        "new_this_month": new_this_month,
        "morgan_connected": morgan_connected
    }

@app.put("/api/admin/users/{user_id}/role")
async def update_user_role(
    user_id: int,
    new_role: str,
    user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update user role (admin only)"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    if new_role not in ["student", "admin"]:
        raise HTTPException(status_code=400, detail="Role must be 'student' or 'admin'")

    target_user = db.query(User).filter(User.id == user_id).first()
    if not target_user:
        raise HTTPException(status_code=404, detail="User not found")

    target_user.role = new_role
    db.commit()

    return {"message": f"User {target_user.email} role updated to {new_role}"}

# --- Admin: System Health ---
@app.get("/api/admin/health")
async def get_system_health(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Get detailed system health (admin only)"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    health_status = {
        "database": {"status": "unknown", "message": ""},
        "vertex_agent": {"status": "unknown", "message": ""},
        "openai_tts": {"status": "unknown", "message": ""},
        "mode": "vertex_ai" if USE_VERTEX_AGENT else "legacy_rag",
        "last_check": datetime.now(timezone.utc).isoformat()
    }

    # Check Database
    try:
        db.execute(text("SELECT 1"))
        health_status["database"] = {"status": "connected", "message": "Database connection OK"}
    except Exception as e:
        health_status["database"] = {"status": "error", "message": str(e)[:100]}

    # Check Vertex AI Agent
    if USE_VERTEX_AGENT:
        health_status["vertex_agent"] = check_agent_health()
    else:
        # Legacy: check Pinecone
        try:
            if PINECONE_API_KEY and PINECONE_INDEX and LEGACY_RAG_AVAILABLE:
                pc_check = Pinecone(api_key=PINECONE_API_KEY)
                idx = pc_check.Index(PINECONE_INDEX)
                stats = idx.describe_index_stats()
                vector_count = stats.get("total_vector_count", 0)
                health_status["vertex_agent"] = {"status": "n/a (legacy mode)", "message": f"Pinecone: {vector_count} vectors"}
            else:
                health_status["vertex_agent"] = {"status": "not_configured", "message": "Legacy mode, keys missing"}
        except Exception as e:
            health_status["vertex_agent"] = {"status": "error", "message": str(e)[:100]}

    # Check OpenAI TTS
    try:
        if OPENAI_API_KEY:
            health_status["openai_tts"] = {"status": "configured", "message": "TTS API key present"}
        else:
            health_status["openai_tts"] = {"status": "not_configured", "message": "TTS unavailable (no OpenAI key)"}
    except Exception as e:
        health_status["openai_tts"] = {"status": "error", "message": str(e)[:100]}

    return health_status

# --- Admin: Course Edit ---
@app.put("/api/curriculum/{code}")
async def update_course(code: str, course: Course, user=Depends(get_current_user)):
    """Update an existing course (admin only)"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admins only")

    arr = json.load(open(CLASSES_FILE, encoding="utf-8"))
    found = False
    for i, c in enumerate(arr):
        if c.get("course_code") == code:
            arr[i] = course.model_dump()
            found = True
            break

    if not found:
        raise HTTPException(status_code=404, detail=f"Course {code} not found")

    json.dump(arr, open(CLASSES_FILE, "w", encoding="utf-8"), indent=2)
    return {"message": f"Course {code} updated", "course": course}

# --- Admin: Knowledge Base Management ---
DATA_SOURCES_DIR = os.path.join(BACKEND_DIR, "data_sources")

@app.get("/api/admin/knowledge-base/files")
async def list_kb_files(user: dict = Depends(get_current_user)):
    """List all knowledge base JSON files (admin only)"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    files = []
    if os.path.exists(DATA_SOURCES_DIR):
        for f in os.listdir(DATA_SOURCES_DIR):
            if f.endswith(".json"):
                filepath = os.path.join(DATA_SOURCES_DIR, f)
                size = os.path.getsize(filepath)
                modified = datetime.fromtimestamp(os.path.getmtime(filepath))
                files.append({
                    "filename": f,
                    "size": size,
                    "modified": modified.isoformat()
                })

    return {"files": sorted(files, key=lambda x: x["filename"])}

@app.get("/api/admin/knowledge-base/search")
async def search_kb_files(q: str, user: dict = Depends(get_current_user)):
    """Search across all knowledge base files (admin only)"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    if not q or len(q) < 2:
        return {"results": []}

    results = []
    search_term = q.lower()

    if os.path.exists(DATA_SOURCES_DIR):
        for filename in os.listdir(DATA_SOURCES_DIR):
            if not filename.endswith(".json"):
                continue

            filepath = os.path.join(DATA_SOURCES_DIR, filename)
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    content = f.read()

                content_lower = content.lower()

                # Find ALL matches in this file
                idx = 0
                match_count = 0
                while True:
                    idx = content_lower.find(search_term, idx)
                    if idx == -1:
                        break

                    match_count += 1

                    # Get context around match (80 chars before and after)
                    start = max(0, idx - 80)
                    end = min(len(content), idx + len(q) + 80)
                    context = content[start:end]

                    # Clean up context (remove newlines for display)
                    context = context.replace('\n', ' ').replace('\r', '')

                    # Find the match in context and highlight it
                    match_start_in_context = idx - start
                    actual_match = content[idx:idx+len(q)]

                    # Build highlighted context
                    highlighted = (
                        context[:match_start_in_context] +
                        f"<mark>{actual_match}</mark>" +
                        context[match_start_in_context + len(q):]
                    )

                    results.append({
                        "filename": filename,
                        "context": "..." + highlighted.strip() + "...",
                        "position": idx,
                        "match_number": match_count
                    })

                    idx += len(q)

                    # Limit matches per file to 10
                    if match_count >= 10:
                        break

            except Exception:
                continue

    # Sort by filename, then position
    results.sort(key=lambda x: (x["filename"], x.get("position", 0)))

    return {"results": results[:50], "total_matches": len(results)}

@app.get("/api/admin/knowledge-base/{filename}")
async def get_kb_file(filename: str, user: dict = Depends(get_current_user)):
    """Get content of a knowledge base file (admin only)"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    if not filename.endswith(".json"):
        raise HTTPException(status_code=400, detail="Only JSON files allowed")

    # Prevent path traversal: strip directory components
    safe_filename = os.path.basename(filename)
    filepath = os.path.join(DATA_SOURCES_DIR, safe_filename)
    if not os.path.realpath(filepath).startswith(os.path.realpath(DATA_SOURCES_DIR)):
        raise HTTPException(status_code=400, detail="Invalid filename")
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")

    try:
        with open(filepath, "r", encoding="utf-8") as f:
            content = json.load(f)
        return {"filename": safe_filename, "content": content}
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"Invalid JSON: {str(e)}")

@app.put("/api/admin/knowledge-base/{filename}")
async def update_kb_file(filename: str, content: dict, user: dict = Depends(get_current_user)):
    """Update a knowledge base file (admin only)"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    if not filename.endswith(".json"):
        raise HTTPException(status_code=400, detail="Only JSON files allowed")

    # Prevent path traversal
    safe_filename = os.path.basename(filename)
    filepath = os.path.join(DATA_SOURCES_DIR, safe_filename)
    if not os.path.realpath(filepath).startswith(os.path.realpath(DATA_SOURCES_DIR)):
        raise HTTPException(status_code=400, detail="Invalid filename")

    # Create backup
    if os.path.exists(filepath):
        backup_path = filepath + ".backup"
        shutil.copy(filepath, backup_path)

    try:
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(content, f, indent=2, ensure_ascii=False)
        return {"message": f"File {filename} updated successfully"}
    except Exception as e:
        # Restore backup on failure
        if os.path.exists(filepath + ".backup"):
            shutil.copy(filepath + ".backup", filepath)
        raise HTTPException(status_code=500, detail=f"Failed to save: {str(e)}")

@app.post("/api/admin/knowledge-base/ingest")
async def trigger_ingestion(user: dict = Depends(get_current_user)):
    """Trigger knowledge base re-ingestion (admin only)"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    try:
        # Legacy Pinecone ingestion removed. Using Vertex AI structured datastore now.
        return {"message": "Ingestion not needed. Using Vertex AI structured datastore (instant updates via admin dashboard)."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ingestion failed: {str(e)}")

# --- Admin: Cloud Knowledge Base (Vertex AI Datastore) ---
from datastore_manager import (
    list_datastore_documents,
    get_document_content,
    upload_document,
    delete_document,
    update_document,
    sync_datastore,
    search_documents as search_cloud_kb,
)

_cloud_kb_cache = {"docs": None, "ts": 0}

@app.get("/api/admin/cloud-kb/documents")
async def list_cloud_kb_docs(user: dict = Depends(get_current_user), refresh: bool = False):
    """List all documents in the Vertex AI Search datastore. Cached for 60s."""
    import time as _t
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    try:
        # Use cached result if fresh (60s TTL) unless forced refresh
        if not refresh and _cloud_kb_cache["docs"] and _t.time() - _cloud_kb_cache["ts"] < 60:
            docs = _cloud_kb_cache["docs"]
            print(f"[CACHE] Cloud KB docs from cache ({len(docs)} docs)")
        else:
            docs = await asyncio.to_thread(list_datastore_documents)
            _cloud_kb_cache["docs"] = docs
            _cloud_kb_cache["ts"] = _t.time()
        return {"documents": docs, "total": len(docs)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to list documents: {e}")

@app.get("/api/admin/cloud-kb/documents/{doc_id}/content")
async def read_cloud_kb_doc(doc_id: str, uri: str = "", user: dict = Depends(get_current_user)):
    """Read content of a document from the structured datastore"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    try:
        content = get_document_content(doc_id)
        return {"content": content, "doc_id": doc_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read document: {e}")

@app.post("/api/admin/cloud-kb/upload")
async def upload_cloud_kb_doc(
    file: UploadFile = File(...),
    user: dict = Depends(get_current_user)
):
    """Upload a new document to the cloud KB"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    allowed_exts = {'txt', 'pdf', 'html', 'csv', 'json'}
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in allowed_exts:
        raise HTTPException(status_code=400, detail=f"Allowed types: {', '.join(allowed_exts)}")

    content = await file.read()
    content_type = file.content_type or "text/plain"

    result = upload_document(file.filename, content, content_type)
    if not result["success"]:
        raise HTTPException(status_code=500, detail=result["message"])
    # Auto-clear cache so chatbot uses fresh data
    cleared = query_cache.clear()
    result["cache_cleared"] = cleared
    return result

@app.put("/api/admin/cloud-kb/documents/{doc_id}")
async def update_cloud_kb_doc(
    doc_id: str,
    request: Request,
    user: dict = Depends(get_current_user)
):
    """Update content of an existing document in the cloud KB"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    body = await request.json()
    content = body.get("content", "")
    if not content:
        raise HTTPException(status_code=400, detail="Content required")

    result = update_document(doc_id, content.encode("utf-8"))
    if not result["success"]:
        raise HTTPException(status_code=500, detail=result["message"])
    # Clear ALL caches + reset ALL ADK sessions so chatbot uses fresh data
    cleared = query_cache.clear()
    # Reset all ADK sessions so no agent reuses stale context
    try:
        from vertex_agent import _session_cache
        session_count = len(_session_cache)
        _session_cache.clear()
    except Exception:
        session_count = 0
    result["cache_cleared"] = cleared
    result["sessions_reset"] = session_count
    return result

@app.delete("/api/admin/cloud-kb/documents/{doc_id}")
async def delete_cloud_kb_doc(doc_id: str, uri: str = "", user: dict = Depends(get_current_user)):
    """Delete a document from the cloud KB"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    result = delete_document(doc_id, uri)
    if not result["success"]:
        raise HTTPException(status_code=500, detail=result["message"])
    # Auto-clear cache so chatbot uses fresh data
    cleared = query_cache.clear()
    result["cache_cleared"] = cleared
    return result

@app.post("/api/admin/cloud-kb/sync")
async def sync_cloud_kb(user: dict = Depends(get_current_user)):
    """Re-sync all GCS documents into the datastore"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    result = sync_datastore()
    if not result["success"]:
        raise HTTPException(status_code=500, detail=result["message"])
    # Auto-clear cache so chatbot uses fresh data
    cleared = query_cache.clear()
    result["cache_cleared"] = cleared
    return result


# ==============================================================================
# CACHE MANAGEMENT ENDPOINTS
# ==============================================================================

@app.get("/api/cache/stats")
async def get_cache_stats_public():
    """Get cache statistics (public, read-only)."""
    stats = query_cache.get_stats()
    return {
        "success": True,
        "cache_stats": stats,
        "cache_type": "multi-tier (L1: in-memory, L2: Redis)"
    }

@app.get("/api/admin/cache/stats")
async def get_cache_stats_admin(user: dict = Depends(get_current_user)):
    """Get cache statistics - admin version with more details."""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    stats = query_cache.get_stats()
    return {
        "success": True,
        "cache_stats": stats
    }

@app.post("/api/admin/cache/clear")
async def clear_cache(user: dict = Depends(get_current_user)):
    """Clear all cached responses"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    cleared_count = query_cache.clear()
    return {
        "success": True,
        "message": f"Cleared {cleared_count} cached items"
    }

@app.get("/api/admin/cloud-kb/search")
async def search_cloud_kb_docs(q: str, user: dict = Depends(get_current_user)):
    """Search across all cloud KB documents"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    if not q or len(q) < 2:
        return {"results": []}
    try:
        results = search_cloud_kb(q)
        return {"results": results, "query": q, "total": len(results)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Search failed: {e}")

# --- Admin: Analytics ---
@app.get("/api/admin/analytics")
async def get_analytics(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Get usage analytics (admin only)"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    from datetime import timedelta
    now = datetime.now(timezone.utc)

    # User signups by day (last 7 days)
    signups_by_day = []
    for i in range(6, -1, -1):
        day = now - timedelta(days=i)
        day_start = day.replace(hour=0, minute=0, second=0, microsecond=0)
        day_end = day_start + timedelta(days=1)
        count = db.query(User).filter(
            User.created_at >= day_start,
            User.created_at < day_end
        ).count()
        signups_by_day.append({
            "date": day_start.strftime("%Y-%m-%d"),
            "day": day_start.strftime("%a"),
            "count": count
        })

    # Ticket stats
    total_tickets = db.query(SupportTicket).count()
    open_tickets = db.query(SupportTicket).filter(SupportTicket.status == "open").count()

    return {
        "signups_by_day": signups_by_day,
        "total_users": db.query(User).count(),
        "total_tickets": total_tickets,
        "open_tickets": open_tickets,
        "timestamp": now.isoformat()
    }

# ==============================================================================
# SUPPORT TICKET ENDPOINTS
# ==============================================================================

@app.get("/api/tickets")
async def list_tickets(status: str = None, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """List tickets - admins see all, users see their own"""
    query = db.query(SupportTicket)
    if user.get("role") != "admin":
        query = query.filter(SupportTicket.user_id == user["user_id"])
    if status and status != "all":
        query = query.filter(SupportTicket.status == status)
    tickets = query.order_by(SupportTicket.created_at.desc()).all()
    return {
        "tickets": [
            {
                "id": t.id,
                "subject": t.subject,
                "category": t.category,
                "description": t.description,
                "status": t.status,
                "priority": t.priority,
                "user_email": db.query(User).filter(User.id == t.user_id).first().email if t.user_id else "Unknown",
                "attachment_name": t.attachment_name,
                "attachment_data": t.attachment_data if t.attachment_data else None,
                "admin_notes": t.admin_notes,
                "created_at": t.created_at.isoformat() if t.created_at else None,
                "updated_at": t.updated_at.isoformat() if t.updated_at else None,
            }
            for t in tickets
        ]
    }

@app.get("/api/tickets/stats/summary")
async def get_ticket_stats(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Get ticket statistics"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    total = db.query(SupportTicket).count()
    open_count = db.query(SupportTicket).filter(SupportTicket.status == "open").count()
    in_progress = db.query(SupportTicket).filter(SupportTicket.status == "in_progress").count()
    resolved = db.query(SupportTicket).filter(SupportTicket.status == "resolved").count()
    return {"total": total, "open": open_count, "in_progress": in_progress, "resolved": resolved}

@app.post("/api/tickets")
async def create_ticket(request: Request, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Create a new support ticket"""
    body = await request.json()
    subject = (body.get("subject", "") or "")[:200]
    description = (body.get("description", "") or "")[:5000]
    category = body.get("category", "other") or "other"
    priority = body.get("priority", "normal") or "normal"
    attachment_data = body.get("attachment_data")
    # Cap base64 attachment at ~7.5MB (10MB file base64-encoded)
    if attachment_data and len(attachment_data) > 10_000_000:
        raise HTTPException(413, "Attachment too large")
    ticket = SupportTicket(
        user_id=user["user_id"],
        subject=subject,
        category=category,
        description=description,
        priority=priority,
        attachment_data=attachment_data,
        attachment_name=(body.get("attachment_name", "") or "")[:255],
    )
    db.add(ticket)
    db.commit()
    db.refresh(ticket)
    return {"success": True, "ticket_id": ticket.id}

@app.put("/api/tickets/{ticket_id}")
async def update_ticket(ticket_id: int, request: Request, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Update ticket status/notes"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    ticket = db.query(SupportTicket).filter(SupportTicket.id == ticket_id).first()
    if not ticket:
        raise HTTPException(status_code=404, detail="Ticket not found")
    body = await request.json()
    if "status" in body:
        ticket.status = body["status"]
        if body["status"] == "resolved":
            ticket.resolved_by = user["user_id"]
            ticket.resolved_at = datetime.now(timezone.utc)
    if "admin_notes" in body:
        ticket.admin_notes = body["admin_notes"]
    db.commit()
    return {"success": True}


# ==============================================================================
# FEEDBACK ENDPOINTS
# ==============================================================================

@app.post("/api/feedback")
async def submit_feedback(request: Request, user: dict = Depends(get_current_user)):
    """Submit feedback on a bot response (helpful/not_helpful/report)."""
    body = await request.json()
    message_text = body.get("message_text", "")
    feedback_type = body.get("feedback_type", "")
    report_details = body.get("report_details", "")
    session_id = body.get("session_id", "default")

    if feedback_type not in ("helpful", "not_helpful", "report"):
        raise HTTPException(status_code=400, detail="Invalid feedback type")

    with SessionLocal() as db:
        fb = Feedback(
            user_id=user.get("user_id"),
            session_id=session_id,
            message_text=message_text[:2000],
            feedback_type=feedback_type,
            report_details=report_details[:1000] if report_details else None,
        )
        db.add(fb)
        db.commit()

    # If "report" (explicit bug report), log as failed query for research.
    # "not_helpful" alone is NOT logged - users thumb-down for many reasons
    # (too verbose, wrong tone, etc.) that don't indicate a KB miss.
    # Only "report" means "this answer is factually wrong or missing info".
    if feedback_type == "report" and message_text:
        try:
            from models import FailedQuery
            with SessionLocal() as db:
                chat = db.query(ChatHistory).filter(
                    ChatHistory.user_id == user.get("user_id"),
                    ChatHistory.bot_response.contains(message_text[:100])
                ).order_by(ChatHistory.timestamp.desc()).first()
                if chat:
                    # Don't duplicate: check if this query was already logged
                    existing = db.query(FailedQuery).filter(
                        FailedQuery.user_query == chat.user_query.strip(),
                        FailedQuery.user_id == user.get("user_id"),
                    ).first()
                    if not existing:
                        entry = FailedQuery(
                            user_query=chat.user_query.strip(),
                            bot_response=chat.bot_response[:1000],
                            user_id=user.get("user_id"),
                            status="new",
                        )
                        db.add(entry)
                        db.commit()
        except Exception:
            pass

    return {"success": True}

@app.get("/api/feedback/stats")
async def get_feedback_stats(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Get feedback statistics"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    total = db.query(Feedback).count()
    helpful = db.query(Feedback).filter(Feedback.feedback_type == "helpful").count()
    not_helpful = db.query(Feedback).filter(Feedback.feedback_type == "not_helpful").count()
    reports = db.query(Feedback).filter(Feedback.feedback_type == "report").count()
    satisfaction_rate = round((helpful / total * 100) if total > 0 else 0, 1)

    # Recent reports
    recent_reports = db.query(Feedback).filter(
        Feedback.feedback_type == "report"
    ).order_by(Feedback.timestamp.desc()).limit(10).all()

    return {
        "total": total,
        "helpful": helpful,
        "not_helpful": not_helpful,
        "reports": reports,
        "satisfaction_rate": satisfaction_rate,
        "recent_reports": [
            {
                "id": r.id,
                "message_preview": (r.message_text[:150] + "...") if r.message_text and len(r.message_text) > 150 else r.message_text,
                "details": r.report_details,
                "timestamp": r.timestamp.isoformat() if r.timestamp else None,
            }
            for r in recent_reports
        ]
    }

@app.get("/api/feedback/all")
async def get_all_feedback(type: str = None, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Get all feedback entries"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    query = db.query(Feedback)
    if type and type != "all":
        query = query.filter(Feedback.feedback_type == type)
    items = query.order_by(Feedback.timestamp.desc()).limit(100).all()
    return {
        "feedback": [
            {
                "id": f.id,
                "user_id": f.user_id,
                "session_id": f.session_id,
                "message_text": f.message_text,
                "feedback_type": f.feedback_type,
                "report_details": f.report_details,
                "timestamp": f.timestamp.isoformat() if f.timestamp else None,
            }
            for f in items
        ]
    }


# ==============================================================================
# AUTO-RESEARCH AGENT ENDPOINTS
# ==============================================================================

from research_agent import run_research_batch, get_research_stats
from models import FailedQuery, KBSuggestion

@app.post("/api/admin/research/run")
async def trigger_research(user: dict = Depends(get_current_user)):
    """Manually trigger a research batch (admin only)."""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    result = await asyncio.to_thread(run_research_batch)
    return result

@app.get("/api/admin/research/stats")
async def research_stats_endpoint(user: dict = Depends(get_current_user)):
    """Get research agent stats for dashboard."""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return get_research_stats()

@app.get("/api/admin/research/suggestions")
async def list_suggestions(status: str = "pending", user: dict = Depends(get_current_user)):
    """List KB suggestions from the research agent."""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    with SessionLocal() as db:
        query = db.query(KBSuggestion)
        if status != "all":
            query = query.filter(KBSuggestion.status == status)
        suggestions = query.order_by(KBSuggestion.created_at.desc()).limit(100).all()
        return {"suggestions": [{
            "id": s.id, "cluster_id": s.cluster_id, "topic": s.topic,
            "representative_query": s.representative_query, "query_count": s.query_count,
            "researched_answer": s.researched_answer,
            "sources": json.loads(s.sources) if s.sources else [],
            "confidence": s.confidence, "suggested_doc_id": s.suggested_doc_id,
            "suggested_content": s.suggested_content, "status": s.status,
            "admin_notes": s.admin_notes,
            "created_at": s.created_at.isoformat() if s.created_at else "",
        } for s in suggestions]}

@app.put("/api/admin/research/suggestions/{suggestion_id}")
async def review_suggestion(suggestion_id: int, request: Request, user: dict = Depends(get_current_user)):
    """Approve, reject, or edit a KB suggestion."""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    body = await request.json()
    action = body.get("action")

    with SessionLocal() as db:
        suggestion = db.query(KBSuggestion).filter(KBSuggestion.id == suggestion_id).first()
        if not suggestion:
            raise HTTPException(status_code=404, detail="Suggestion not found")

        if action == "approve":
            suggestion.status = "approved"
            suggestion.reviewed_by = user.get("user_id")
            suggestion.reviewed_at = datetime.now(timezone.utc)
        elif action == "reject":
            suggestion.status = "rejected"
            suggestion.admin_notes = body.get("notes", "")
            suggestion.reviewed_by = user.get("user_id")
            suggestion.reviewed_at = datetime.now(timezone.utc)
        elif action == "edit":
            if "content" in body:
                suggestion.suggested_content = body["content"]
            if "doc_id" in body:
                suggestion.suggested_doc_id = body["doc_id"]
            if "notes" in body:
                suggestion.admin_notes = body["notes"]

        db.commit()
    return {"success": True}

@app.post("/api/admin/research/suggestions/{suggestion_id}/push")
async def push_suggestion(suggestion_id: int, user: dict = Depends(get_current_user)):
    """Push an approved suggestion to the live KB datastore."""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    with SessionLocal() as db:
        suggestion = db.query(KBSuggestion).filter(
            KBSuggestion.id == suggestion_id,
            KBSuggestion.status == "approved"
        ).first()
        if not suggestion:
            raise HTTPException(status_code=404, detail="Approved suggestion not found")

        doc_id = suggestion.suggested_doc_id
        content = suggestion.suggested_content
        if not doc_id or not content:
            raise HTTPException(status_code=400, detail="Missing doc_id or content")

        # Check if doc exists -> append; otherwise -> create
        existing = get_document_content(doc_id)
        if existing and not existing.startswith("Error"):
            merged = existing.rstrip() + "\n\n" + content
            result = update_document(doc_id, merged.encode("utf-8"))
        else:
            result = upload_document(f"{doc_id}.txt", content.encode("utf-8"))

        if result["success"]:
            suggestion.status = "pushed"
            db.commit()
            query_cache.clear()
            try:
                from vertex_agent import _session_cache
                _session_cache.clear()
            except Exception:
                pass
            return {"success": True, "message": f"Pushed to KB as {doc_id}"}
        else:
            raise HTTPException(status_code=500, detail=result["message"])

@app.get("/api/admin/research/failed-queries")
async def list_failed_queries(status: str = "all", user: dict = Depends(get_current_user)):
    """List raw failed queries for transparency."""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    with SessionLocal() as db:
        query = db.query(FailedQuery)
        if status != "all":
            query = query.filter(FailedQuery.status == status)
        queries = query.order_by(FailedQuery.created_at.desc()).limit(200).all()
        return {"queries": [{
            "id": q.id, "user_query": q.user_query, "bot_response": q.bot_response[:200],
            "cluster_id": q.cluster_id, "status": q.status,
            "created_at": q.created_at.isoformat() if q.created_at else "",
        } for q in queries]}

@app.post("/api/internal/research/run")
async def internal_research_trigger(request: Request):
    """Triggered by Cloud Scheduler daily at 2am. Auth via shared secret."""
    secret = request.headers.get("X-Research-Secret", "")
    expected = os.getenv("RESEARCH_SECRET", "")
    if not expected or secret != expected:
        raise HTTPException(status_code=403, detail="Invalid research secret")
    result = await asyncio.to_thread(run_research_batch)
    return result


@app.post("/api/internal/memory/consolidate")
async def internal_memory_consolidate(request: Request):
    """Triggered by Cloud Scheduler daily at 3am. Consolidates conversations into long-term user memories."""
    secret = request.headers.get("X-Research-Secret", "")
    expected = os.getenv("RESEARCH_SECRET", "")
    if not expected or secret != expected:
        raise HTTPException(status_code=403, detail="Invalid research secret")
    from services.memory_service import consolidate_user_memories
    result = await asyncio.to_thread(consolidate_user_memories, 24)
    return result


@app.post("/api/internal/canvas/sync")
async def internal_canvas_sync(request: Request):
    """Triggered by Cloud Scheduler daily at 4am. Refreshes Canvas data for all synced users.
    Requires canvas_client.refresh_canvas_data() to be implemented."""
    secret = request.headers.get("X-Research-Secret", "")
    expected = os.getenv("RESEARCH_SECRET", "")
    if not expected or secret != expected:
        raise HTTPException(status_code=403, detail="Invalid research secret")

    # Canvas uses LDAP session auth. Cannot auto-refresh without storing credentials.
    # This endpoint reports stale records so admins know which students have old data.
    from models import CanvasStudentData
    from datetime import timedelta

    db = SessionLocal()
    try:
        canvas_records = db.query(CanvasStudentData).all()
        if not canvas_records:
            return {"status": "no_canvas_users", "total": 0}

        stale_cutoff = datetime.utcnow() - timedelta(days=3)
        stale = [r for r in canvas_records if r.synced_at and r.synced_at < stale_cutoff]
        fresh = [r for r in canvas_records if r.synced_at and r.synced_at >= stale_cutoff]

        return {
            "status": "report",
            "note": "Canvas uses LDAP auth. Cannot auto-refresh. Students re-sync manually in Profile.",
            "total_users": len(canvas_records),
            "fresh_last_3d": len(fresh),
            "stale_over_3d": len(stale),
        }
    finally:
        db.close()


@app.post("/api/internal/degreeworks/sync")
async def internal_degreeworks_sync(request: Request):
    """Triggered by Cloud Scheduler monthly (1st of month at 5am). Refreshes DegreeWorks data.
    Requires banner_scraper.client.refresh_degreeworks_data() to be implemented."""
    secret = request.headers.get("X-Research-Secret", "")
    expected = os.getenv("RESEARCH_SECRET", "")
    if not expected or secret != expected:
        raise HTTPException(status_code=403, detail="Invalid research secret")

    # DegreeWorks uses CAS session auth (no API tokens). Cannot auto-refresh
    # without student credentials. This endpoint reports stale records for admin awareness.
    from models import DegreeWorksData
    from datetime import timedelta

    db = SessionLocal()
    try:
        dw_records = db.query(DegreeWorksData).all()
        if not dw_records:
            return {"status": "no_degreeworks_users", "total": 0}

        stale_cutoff = datetime.utcnow() - timedelta(days=30)
        stale = [r for r in dw_records if r.synced_at and r.synced_at < stale_cutoff]
        fresh = [r for r in dw_records if r.synced_at and r.synced_at >= stale_cutoff]

        return {
            "status": "report",
            "note": "DegreeWorks requires CAS login. Cannot auto-refresh. Students must re-sync manually.",
            "total_users": len(dw_records),
            "fresh_last_30d": len(fresh),
            "stale_over_30d": len(stale),
        }
    finally:
        db.close()


# ==============================================================================
# CLOUD KB STATS ENDPOINT
# ==============================================================================

@app.get("/api/admin/cloud-kb/stats")
async def get_cloud_kb_stats(user: dict = Depends(get_current_user)):
    """Get cloud KB statistics - doc count, total size, last modified"""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    try:
        docs = list_datastore_documents()
        total_size = sum(d.get("size", 0) for d in docs)
        last_modified = max((d.get("modified", "") for d in docs), default="") if docs else ""
        return {
            "total_documents": len(docs),
            "total_size": total_size,
            "last_modified": last_modified,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get stats: {e}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=5000, reload=True)